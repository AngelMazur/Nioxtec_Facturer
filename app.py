"""
Nioxtec Facturer
=================

This Flask application implements a very small invoice/proforma
management system designed for sole traders and small businesses.
It exposes a REST‑style API for creating clients and invoices as
well as an HTML front‑end that allows you to fill in forms and
generate PDFs.  Everything is kept intentionally simple so that
non‑programmers can understand the overall flow.

The core ideas implemented here follow widely accepted
recommendations for invoice management.  For example, the EU
directive on invoicing requires that each invoice include a unique
sequential number, the date of issue, the supplier's and customer's
VAT numbers, the full names and addresses of both parties and the
detailed description of goods or services provided【816995957740757†L97-L115】.  In
addition, the database schema uses one‑to‑many relationships
between invoices and their line items so that each invoice can
contain multiple entries without duplicating client data【816995957740757†L130-L151】.

The application also relies on WeasyPrint to render HTML and CSS
templates as high quality PDFs.  WeasyPrint is a free and open
source library that acts like a headless browser – it supports
modern CSS features such as flexbox and grid and integrates
seamlessly with Jinja templates【239017722105616†L76-L83】.  This makes it ideal
for generating invoices, reports and other documents【239017722105616†L83-L94】.
"""

from datetime import datetime, timedelta
import os
import time
import re
import unicodedata
from docx import Document
from flask import Flask, jsonify, request, render_template, send_file, abort, url_for, Response
from flask_sqlalchemy import SQLAlchemy
from flask_compress import Compress
from flask_cors import CORS
from flask_talisman import Talisman
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_jwt_extended import (
    JWTManager, create_access_token, jwt_required, get_jwt_identity,
    set_access_cookies, unset_jwt_cookies, verify_jwt_in_request
)
from sqlalchemy import inspect, text
from sqlalchemy.exc import IntegrityError
from dotenv import load_dotenv
from types import SimpleNamespace
try:
    import sentry_sdk  # type: ignore
    from sentry_sdk.integrations.flask import FlaskIntegration  # type: ignore
except Exception:
    sentry_sdk = None  # type: ignore
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from typing import Tuple
from werkzeug.middleware.proxy_fix import ProxyFix
import logging
import traceback
from pydantic import ValidationError as PydValidationError

from schemas.api import LoginRequest, ClientCreateRequest, InvoiceCreateRequest
from openapi import get_openapi_spec

# Cargar variables de entorno desde .env si existe
load_dotenv()

# Sentry opcional (activado solo si SENTRY_DSN está definido)
try:
    _sentry_dsn = os.getenv('SENTRY_DSN')
    if _sentry_dsn and sentry_sdk is not None:
        sentry_sdk.init(
            dsn=_sentry_dsn,
            integrations=[FlaskIntegration()],
            traces_sample_rate=float(os.getenv('SENTRY_TRACES_SAMPLE_RATE', '0.0')),
            environment=os.getenv('APP_ENV', 'production'),
            release=os.getenv('APP_VERSION') or None,
        )
except Exception:
    # No bloquear el arranque si Sentry falla
    pass

# Motor PDF unificado: solo wkhtmltopdf para consistencia dev/prod
# Eliminamos WeasyPrint para evitar incompatibilidades entre entornos
use_weasyprint = False

try:
    import pdfkit  # type: ignore
except Exception:
    pdfkit = None

def _resolve_pdfkit_configuration():
    """Return a pdfkit configuration resolving wkhtmltopdf path on Windows if needed."""
    if pdfkit is None:
        return None
    try:
        from shutil import which
        wkhtml_path = os.getenv('WKHTMLTOPDF_PATH') or which('wkhtmltopdf')
        if not wkhtml_path:
            # Common Windows install paths
            candidates = [
                r"C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe",
                r"C:\\Program Files (x86)\\wkhtmltopdf\\bin\\wkhtmltopdf.exe",
            ]
            for p in candidates:
                if os.path.isfile(p):
                    wkhtml_path = p
                    break
        return pdfkit.configuration(wkhtmltopdf=wkhtml_path) if wkhtml_path else pdfkit.configuration()
    except Exception:
        return None

# Fallback PDF generator (pure Python) when WeasyPrint/wkhtmltopdf are unavailable
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import mm
    reportlab_available = True
except Exception:
    reportlab_available = False


# -------------------------------------------------------------
#  config helpers

def _company_from_env() -> SimpleNamespace:
    """Build a company-like object from environment variables.

    Allows showing real company data in PDFs without requiring a DB row.

    Supported env vars:
      COMPANY_NAME, COMPANY_CIF, COMPANY_ADDRESS, COMPANY_CITY,
      COMPANY_PROVINCE, COMPANY_EMAIL, COMPANY_PHONE, COMPANY_IBAN,
      COMPANY_WEBSITE
    """
    return SimpleNamespace(
        name=os.getenv('COMPANY_NAME', 'Mi Empresa'),
        cif=os.getenv('COMPANY_CIF', 'A00000000'),
        address=os.getenv('COMPANY_ADDRESS', 'Dirección de ejemplo'),
        city=os.getenv('COMPANY_CITY', ''),
        province=os.getenv('COMPANY_PROVINCE', ''),
        email=os.getenv('COMPANY_EMAIL', 'info@example.com'),
        phone=os.getenv('COMPANY_PHONE', '000 000 000'),
        iban=os.getenv('COMPANY_IBAN', ''),
        website=os.getenv('COMPANY_WEBSITE', ''),
    )


def _compose_company_address(company) -> str:
    """Compose a single address line without repeating city/province.

    - Starts from company.address
    - Appends city and province only if they are not already present
      and not equal to each other (case-insensitive)
    """
    base = (getattr(company, 'address', '') or '').strip(' ,\n\t')
    addr_lower = base.lower()
    parts = [base] if base else []

    def add_if_absent(value: str):
        v = (value or '').strip()
        if not v:
            return
        if v.lower() not in addr_lower and all(p.lower() != v.lower() for p in parts):
            parts.append(v)

    add_if_absent(getattr(company, 'city', ''))
    # Avoid duplicating province if same as city
    province = getattr(company, 'province', '') or ''
    if province.strip().lower() != (getattr(company, 'city', '') or '').strip().lower():
        add_if_absent(province)

    return ', '.join([p for p in parts if p])
import io
from pathlib import Path
from uuid import uuid4
import json

# -----------------------------------------------------------------------------
# Flask configuration
#
# For a small project like this SQLite is perfectly adequate.  It stores
# all your data in a single file (app.db) and requires no external server.
# When you grow beyond a single user you can switch SQLALCHEMY_DATABASE_URI
# to a PostgreSQL connection string without changing your code.
app = Flask(__name__, instance_relative_config=True)
# Limitar tamaño de subida global (20 MB)
app.config['MAX_CONTENT_LENGTH'] = int(os.getenv('MAX_CONTENT_LENGTH_MB', '20')) * 1024 * 1024
# Detrás de un proxy/túnel (Cloudflare/Nginx), respeta cabeceras X-Forwarded-*
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_port=1, x_prefix=1)  # type: ignore

# Compresión de respuestas (JSON/HTML) para reducir latencia en listados/reportes
app.config.setdefault('COMPRESS_MIMETYPES', ['application/json', 'text/html'])
app.config.setdefault('COMPRESS_LEVEL', 6)
app.config.setdefault('COMPRESS_MIN_SIZE', 500)
Compress(app)

# Logging estructurado opcional (JSON) controlado por ENV JSON_LOGS=true
class _JsonFormatter(logging.Formatter):
    def format(self, record):
        payload = {
            'level': record.levelname,
            'name': record.name,
            'message': record.getMessage(),
            'time': datetime.utcnow().isoformat() + 'Z',
        }
        if record.exc_info:
            payload['exc_info'] = ''.join(traceback.format_exception(*record.exc_info))
        return json.dumps(payload, ensure_ascii=False)

if os.getenv('JSON_LOGS', 'false').lower() in ('1', 'true', 'yes'):
    try:
        handler = logging.StreamHandler()
        handler.setFormatter(_JsonFormatter())
        root = logging.getLogger()
        root.handlers = [handler]
        root.setLevel(logging.INFO)
    except Exception:
        pass

# Filtro personalizado para saltos de línea en PDFs
@app.template_filter('nl2br')
def nl2br_filter(text):
    """Convierte saltos de línea (\n) a <br> tags más compatibles con wkhtmltopdf."""
    if not text:
        return ''
    from markupsafe import Markup
    # Usar <br /> con espacio antes para mejor compatibilidad con wkhtmltopdf
    return Markup(text.replace('\n', '<br />\n'))

@app.template_filter('nl2div')
def nl2div_filter(text):
    """Convierte saltos de línea (\n) a divs separados - más compatible con wkhtmltopdf."""
    if not text:
        return ''
    from markupsafe import Markup
    lines = text.split('\n')
    if len(lines) <= 1:
        return Markup(text)
    
    html_lines = []
    for i, line in enumerate(lines):
        if i == 0:
            html_lines.append(line.strip())
        else:
            html_lines.append(f'<div style="margin:0; padding:0; line-height:1.2;">{line.strip()}</div>')
    
    return Markup(''.join(html_lines))

@app.template_filter('nl2p')
def nl2p_filter(text):
    """Convierte saltos de línea (\n) a párrafos <p> - MÁS ROBUSTO con wkhtmltopdf."""
    if not text:
        return ''
    from markupsafe import Markup
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if len(lines) <= 1:
        return Markup(f'<p style="margin:2px 0; padding:0; line-height:1.3;">{text.strip()}</p>')
    
    html_lines = []
    for line in lines:
        if line:  # Solo líneas no vacías
            html_lines.append(f'<p style="margin:2px 0; padding:0; line-height:1.3;">{line}</p>')
    
    return Markup(''.join(html_lines))

@app.template_filter('nl2blocks')
def nl2blocks_filter(text):
    """Solución #1 para wkhtmltopdf 0.12.6: TODAS las líneas en bloques uniformes.
    Evita mezclar texto 'suelto' con bloques - genera bloques homogéneos.
    """
    if not text:
        return ''
    from markupsafe import Markup
    lines = text.split('\n')
    
    # TODAS las líneas van en bloques uniformes (incluida la primera)
    html_lines = []
    for line in lines:
        # Mantener líneas vacías como espacios para preservar el formato
        if line.strip() == '':
            html_lines.append('<span style="display:block; height:1.2em;">&nbsp;</span>')
        else:
            html_lines.append(f'<span style="display:block; margin:0; padding:0; line-height:1.2;">{line.strip()}</span>')
    
    return Markup(''.join(html_lines))

# CORS configurable (por defecto permite localhost dev y nginx)
cors_origins_str = os.getenv('CORS_ORIGINS', 'http://localhost:5173,http://127.0.0.1:5173,http://localhost:8080,http://127.0.0.1:8080')
allowed_origins = [o.strip() for o in cors_origins_str.split(',') if o.strip()]
# Permite inyectar orígenes extra por ENV para despliegues (ej.: https://app.nioxtec.es)
extra_origins_str = os.getenv('EXTRA_ORIGINS') or os.getenv('APP_ORIGIN') or os.getenv('PUBLIC_APP_ORIGIN')
if extra_origins_str:
    for o in [x.strip() for x in extra_origins_str.split(',') if x.strip()]:
        if o not in allowed_origins:
            allowed_origins.append(o)
# En APIs públicas detrás de dominio distinto, necesitamos reflejar el origen y
# responder correctamente a preflights. También habilitamos Authorization.
CORS(
    app,
    resources={r"/*": {"origins": allowed_origins}},
    supports_credentials=True,
    allow_headers=["Content-Type", "Authorization"],
    expose_headers=["Content-Type", "Authorization"],
    methods=["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS"],
    max_age=86400,
)

# Asegurar cabeceras CORS incluso si algún middleware/servidor intermedio
# no propaga correctamente la respuesta de Flask-CORS
@app.after_request
def _add_cors_headers(response: Response) -> Response:
    try:
        origin = request.headers.get('Origin')
        if origin and origin in allowed_origins:
            response.headers['Access-Control-Allow-Origin'] = origin
            response.headers['Vary'] = 'Origin'
            response.headers['Access-Control-Allow-Credentials'] = 'true'
            response.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
            response.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, PATCH, DELETE, OPTIONS'
    except Exception:
        pass
    return response

# JWT
app.config['JWT_SECRET_KEY'] = os.getenv('JWT_SECRET_KEY', 'change-this-secret')
# Flag de debug (usado también para relajar X-Frame-Options solo en local)
DEBUG_MODE = os.getenv('FLASK_DEBUG', 'true').lower() in ('1','true','yes')

# Permitir token en query (?token=...) solo si está habilitado explícitamente.
# En desarrollo: permitido por defecto; en producción: deshabilitado por defecto.
allow_query_token = os.getenv(
    'ALLOW_QUERY_TOKEN', 'true' if DEBUG_MODE else 'false'
).lower() in ('1', 'true', 'yes')

# Opción para desactivar cookies JWT completamente (mejor interoperabilidad móvil)
disable_jwt_cookies = os.getenv('JWT_DISABLE_COOKIES', 'false').lower() in ('1','true','yes')

# Ubicaciones válidas del JWT
jwt_locations = ['headers']
if not disable_jwt_cookies:
    jwt_locations.append('cookies')
if allow_query_token:
    jwt_locations.append('query_string')
app.config['JWT_TOKEN_LOCATION'] = jwt_locations
app.config['JWT_QUERY_STRING_NAME'] = 'token'
# Cookies JWT: seguras en prod, flexibles en dev
app.config['JWT_COOKIE_SECURE'] = not DEBUG_MODE
app.config['JWT_COOKIE_SAMESITE'] = os.getenv('JWT_COOKIE_SAMESITE', 'Lax' if DEBUG_MODE else 'None')
app.config['JWT_COOKIE_DOMAIN'] = os.getenv('JWT_COOKIE_DOMAIN')  # ej.: api.nioxtec.es o .nioxtec.es
app.config['JWT_COOKIE_CSRF_PROTECT'] = False
# Rechazar arranque en prod sin secreto adecuado
if not DEBUG_MODE:
    if not os.getenv('JWT_SECRET_KEY') or app.config['JWT_SECRET_KEY'] == 'change-this-secret':
        raise RuntimeError('JWT_SECRET_KEY debe definirse en producción')
jwt = JWTManager(app)
# Ampliar la vida del token para mantener sesión hasta cerrar (30 días)
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(days=30)
# Asegurar carpeta de instancia (para SQLite por defecto)
os.makedirs(app.instance_path, exist_ok=True)

# Base de datos configurable vía env; por defecto usa instance/app.db
database_url = os.getenv('DATABASE_URL') or f"sqlite:///{os.path.join(app.instance_path, 'app.db')}"
app.config['SQLALCHEMY_DATABASE_URI'] = database_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# -----------------------------------------------------------------------------
# Security hardening (Talisman + Rate Limiting)
# -----------------------------------------------------------------------------

enable_talisman = os.getenv('ENABLE_TALISMAN', 'true').lower() in ('1', 'true', 'yes')
if enable_talisman:
    # CSP mínimo: permite este origen, inline para plantillas simples y blobs para descargas
    # En desarrollo, permitir cargar imágenes desde localhost:5001 (backend)
    img_sources = ["'self'", 'data:', 'blob:']
    if DEBUG_MODE:
        img_sources.extend(['http://localhost:5001', 'http://127.0.0.1:5001'])
    
    csp = {
        'default-src': ["'self'"],
        'script-src': ["'self'", "'unsafe-inline'", 'https://cdn.jsdelivr.net'],
        'style-src': ["'self'", "'unsafe-inline'", 'https://fonts.googleapis.com'],
        'font-src': ["'self'", 'data:', 'https://fonts.gstatic.com'],
        'img-src': img_sources,
            # Permitir previsualización de PDFs en iframe como blob:
            'frame-src': ["'self'", 'blob:'],
        # Permitimos conexiones desde los orígenes del frontend además de self
        'connect-src': ["'self'", *allowed_origins],
        'frame-ancestors': ["'self'"],
    }
    talisman = Talisman(
        app,
        content_security_policy=csp,
        force_https=bool(os.getenv('FORCE_HTTPS', 'true').lower() in ('1','true','yes')),
        strict_transport_security=True,
        strict_transport_security_max_age=31536000,
    )
    # Permitir iframes/preview en local: X-Frame-Options SAMEORIGIN ya es suficiente; no cambiar en prod
    if DEBUG_MODE:
        @app.after_request
        def _relax_xfo(resp: Response) -> Response:
            # Asegurar que blob: en iframe no se bloquee por header conflictivo
            # No se envía en prod (controlado por DEBUG_MODE)
            resp.headers['X-Frame-Options'] = 'SAMEORIGIN'
            return resp

# Rate limiting por IP (puede migrarse a Redis en prod con STORAGE_URI)
limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    default_limits=["1000 per day", "300 per hour"],
    storage_uri=os.getenv('LIMITER_STORAGE_URI', "memory://"),
)

# Directory where generated PDFs will be saved.  This makes it easy for the
# React front‑end to offer downloadable links.  The folder is created on
# startup if it doesn't exist.
DOWNLOAD_FOLDER = os.path.join(app.root_path, 'downloads')
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)

# Asegurar carpeta static para logo PDF
STATIC_FOLDER = os.path.join(app.root_path, 'static')
os.makedirs(STATIC_FOLDER, exist_ok=True)

# Carpeta base para subidas (documentos/imágenes de clientes)
UPLOADS_ROOT = os.path.join(app.instance_path, 'uploads')
os.makedirs(UPLOADS_ROOT, exist_ok=True)


# -------------------------------------
# Error handling (uniform JSON)
# -------------------------------------

@app.errorhandler(400)
def handle_400(err):
    return jsonify({"error": getattr(err, 'description', 'bad request'), "code": 400}), 400

@app.errorhandler(401)
def handle_401(err):
    return jsonify({"error": getattr(err, 'description', 'unauthorized'), "code": 401}), 401

@app.errorhandler(429)
def handle_429(err):
    return jsonify({"error": "rate limit exceeded", "code": 429}), 429

@app.errorhandler(404)
def handle_404(err):
    return jsonify({"error": "not found", "code": 404}), 404

@app.errorhandler(500)
def handle_500(err):
    return jsonify({"error": 'internal error', "code": 500}), 500


# -----------------------------------------------------------------------------
# Database models
#
# Each class below corresponds to a table in the database.  SQLAlchemy
# automatically creates foreign key relationships for us when we reference
# another model.

class CompanyConfig(db.Model):
    """Stores fixed company information printed on every document."""
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(128), nullable=False)
    cif = db.Column(db.String(32), nullable=False)
    address = db.Column(db.String(256), nullable=False)
    city = db.Column(db.String(128))
    province = db.Column(db.String(128))
    email = db.Column(db.String(128), nullable=False)
    phone = db.Column(db.String(64), nullable=False)
    iban = db.Column(db.String(64))
    website = db.Column(db.String(128))


class Client(db.Model):
    """Represents a customer or client."""
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(128), nullable=False)
    cif = db.Column(db.String(32), nullable=False)
    address = db.Column(db.String(256), nullable=False)
    email = db.Column(db.String(128), nullable=False)
    phone = db.Column(db.String(64), nullable=False)
    iban = db.Column(db.String(64))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class Invoice(db.Model):
    """Represents both invoices and proformas."""
    id = db.Column(db.Integer, primary_key=True)
    number = db.Column(db.String(32), unique=True, nullable=False, index=True)
    date = db.Column(db.Date, nullable=False, index=True)
    type = db.Column(db.String(16), nullable=False)  # 'factura' or 'proforma'
    client_id = db.Column(db.Integer, db.ForeignKey('client.id'), nullable=False)
    notes = db.Column(db.Text)
    payment_method = db.Column(db.String(32))  # bizum | efectivo | transferencia
    total = db.Column(db.Float, default=0.0)
    tax_total = db.Column(db.Float, default=0.0)
    paid = db.Column(db.Boolean, default=False, nullable=False, index=True)
    client = db.relationship('Client', backref=db.backref('invoices', lazy=True))


class ClientDocument(db.Model):
    """Archivos asociados a un cliente (PDFs e Imágenes)."""
    id = db.Column(db.Integer, primary_key=True)
    client_id = db.Column(db.Integer, db.ForeignKey('client.id'), nullable=False, index=True)
    # 'document' (PDF) | 'image' (jpg/png/webp)
    category = db.Column(db.String(16), nullable=False)
    filename = db.Column(db.String(255), nullable=False)  # nombre original
    stored_path = db.Column(db.String(512), nullable=False)  # ruta relativa bajo UPLOADS_ROOT
    content_type = db.Column(db.String(128), nullable=False)
    size_bytes = db.Column(db.Integer, default=0)
    uploaded_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    client = db.relationship('Client', backref=db.backref('documents', lazy=True))


class InvoiceItem(db.Model):
    """Line items that belong to an invoice."""
    id = db.Column(db.Integer, primary_key=True)
    invoice_id = db.Column(db.Integer, db.ForeignKey('invoice.id'), nullable=False)
    product_id = db.Column(db.Integer, db.ForeignKey('product.id'), nullable=True, index=True)
    description = db.Column(db.String(512), nullable=False)
    units = db.Column(db.Integer, nullable=False)
    unit_price = db.Column(db.Float, nullable=False)
    tax_rate = db.Column(db.Float, nullable=False)  # as percentage, e.g. 21 for 21%
    subtotal = db.Column(db.Float, nullable=False)
    total = db.Column(db.Float, nullable=False)
    invoice = db.relationship('Invoice', backref=db.backref('items', lazy=True))


class Product(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    category = db.Column(db.String(64), nullable=False, index=True)
    model = db.Column(db.String(128), nullable=False, index=True)
    sku = db.Column(db.String(64), unique=True)
    stock_qty = db.Column(db.Integer, default=0)
    price_net = db.Column(db.Float, nullable=False, default=0.0)
    tax_rate = db.Column(db.Float, default=21.0)
    features = db.Column(db.JSON)
    images = db.Column(db.JSON, default=list)  # Lista de imágenes del producto
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    # Soft-delete / archive flag
    is_active = db.Column(db.Boolean, default=True, index=True)


class StockMovement(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    product_id = db.Column(db.Integer, db.ForeignKey('product.id'), nullable=False, index=True)
    qty = db.Column(db.Integer, nullable=False)
    type = db.Column(db.String(16), nullable=False)  # 'sale' | 'manual' | 'adjust'
    invoice_id = db.Column(db.Integer, db.ForeignKey('invoice.id'))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class User(db.Model):
    """Basic user for authentication."""
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False, index=True)
    password_hash = db.Column(db.String(256), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)


class DocumentSequence(db.Model):
    """Stores the last issued number per document type and period (year, month)."""
    id = db.Column(db.Integer, primary_key=True)
    doc_type = db.Column(db.String(16), nullable=False, index=True)
    year = db.Column(db.Integer, nullable=False, default=0, index=True)
    month = db.Column(db.Integer, nullable=False, default=0, index=True)
    last_number = db.Column(db.Integer, nullable=False, default=0)


class Expense(db.Model):
    """Represents an expense record."""
    id = db.Column(db.Integer, primary_key=True)
    date = db.Column(db.Date, nullable=False, index=True)
    category = db.Column(db.String(64), nullable=False, index=True)
    description = db.Column(db.String(256), nullable=False)
    supplier = db.Column(db.String(128), nullable=False)
    base_amount = db.Column(db.Float, nullable=False)
    tax_rate = db.Column(db.Float, default=21.0)
    total = db.Column(db.Float, nullable=False)
    paid = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


# -----------------------------------------------------------------------------
# Helper functions

def calculate_totals(items):
    """Compute the subtotal, tax amount and total for a list of items."""
    subtotal_sum = sum(item['units'] * item['unit_price'] for item in items)
    tax_sum = sum(item['units'] * item['unit_price'] * (item['tax_rate'] / 100) for item in items)
    total_sum = subtotal_sum + tax_sum
    return subtotal_sum, tax_sum, total_sum


def _format_number_for_type(doc_type: str, sequence_number: int, year: int, month: int) -> str:
    """Return formatted number FAAMM### or PAAMM### depending on type.

    - Prefix: 'F' para factura, 'P' para proforma
    - AAMM: año y mes en dos dígitos cada uno (YYMM)
    - ###: contador con tres dígitos iniciado en 001 cada mes
    """
    prefix = 'F' if doc_type == 'factura' else 'P'
    yy = year % 100
    mm = month % 100
    return f"{prefix}{yy:02d}{mm:02d}{sequence_number:03d}"


def _next_sequence_atomic(doc_type: str, at_date: datetime = None) -> str:
    """Atomically increment and return the next formatted number for a given type and current year/month."""
    at = at_date or datetime.utcnow()
    y, m = at.year, at.month
    # Locking approach: in SQLite rely on transaction; in Postgres, FOR UPDATE ideal (SQLAlchemy hint used).
    seq = (DocumentSequence.query
           .filter_by(doc_type=doc_type, year=y, month=m)
           .with_for_update(nowait=False)
           .first())
    if not seq:
        seq = DocumentSequence(doc_type=doc_type, year=y, month=m, last_number=0)
        db.session.add(seq)
        db.session.flush()
    seq.last_number += 1
    db.session.flush()
    return _format_number_for_type(doc_type, seq.last_number, y, m)


def _generate_pdf_fallback(invoice: 'Invoice', client: 'Client', company: 'CompanyConfig', items) -> bytes:
    """Generate a very simple PDF using ReportLab as a last-resort fallback.
    Avoids external binaries. Intended only when WeasyPrint/pdfkit are unavailable.
    """
    if not reportlab_available:
        abort(500, description='No hay motor PDF disponible. Instale wkhtmltopdf o WeasyPrint.')
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    y = height - 20 * mm
    c.setFont("Helvetica-Bold", 14)
    c.drawString(20 * mm, y, f"{company.name if company and getattr(company, 'name', None) else 'Empresa'}")
    y -= 8 * mm
    c.setFont("Helvetica", 10)
    if company:
        c.drawString(20 * mm, y, f"CIF: {company.cif or ''}  Tel: {company.phone or ''}")
        y -= 6 * mm
        c.drawString(20 * mm, y, f"{company.address or ''}")
        y -= 10 * mm

    c.setFont("Helvetica-Bold", 12)
    c.drawString(20 * mm, y, f"{invoice.type.title()} {invoice.number}")
    y -= 6 * mm
    c.setFont("Helvetica", 10)
    c.drawString(20 * mm, y, f"Fecha: {invoice.date.isoformat()}")
    y -= 10 * mm

    c.setFont("Helvetica-Bold", 11)
    c.drawString(20 * mm, y, "Cliente")
    y -= 6 * mm
    c.setFont("Helvetica", 10)
    c.drawString(20 * mm, y, f"{client.name}  CIF: {client.cif}")
    y -= 6 * mm
    c.drawString(20 * mm, y, f"{client.address}")
    y -= 10 * mm

    c.setFont("Helvetica-Bold", 10)
    c.drawString(20 * mm, y, "Descripción")
    c.drawString(110 * mm, y, "Unidades")
    c.drawString(140 * mm, y, "Precio")
    c.drawString(170 * mm - 10 * mm, y, "IVA%")
    y -= 5 * mm
    c.line(20 * mm, y, 190 * mm, y)
    y -= 5 * mm
    c.setFont("Helvetica", 10)
    subtotal_sum, tax_sum, total_sum = calculate_totals([
        {
            'units': it.units,
            'unit_price': it.unit_price,
            'tax_rate': it.tax_rate,
        } for it in items
    ])
    for it in items:
        if y < 30 * mm:
            c.showPage()
            y = height - 20 * mm
            c.setFont("Helvetica", 10)
        
        # Manejar saltos de línea en la descripción
        description_lines = (it.description or '').split('\n')
        first_line = True
        
        for line in description_lines:
            if y < 30 * mm:
                c.showPage()
                y = height - 20 * mm
                c.setFont("Helvetica", 10)
            
            # Solo mostrar datos numéricos en la primera línea
            if first_line:
                c.drawString(20 * mm, y, line[:80])
                c.drawRightString(135 * mm, y, f"{it.units}")
                c.drawRightString(165 * mm, y, f"{it.unit_price:.2f} €")
                c.drawRightString(190 * mm, y, f"{it.tax_rate:.2f}")
                first_line = False
            else:
                # Líneas adicionales solo con texto de descripción
                c.drawString(20 * mm, y, line[:80])
            
            y -= 6 * mm

    y -= 6 * mm
    c.line(120 * mm, y, 190 * mm, y)
    y -= 8 * mm
    c.setFont("Helvetica-Bold", 11)
    c.drawRightString(165 * mm, y, "Subtotal:")
    c.drawRightString(190 * mm, y, f"{subtotal_sum:.2f} €")
    y -= 6 * mm
    c.setFont("Helvetica", 11)
    c.drawRightString(165 * mm, y, "Impuestos:")
    c.drawRightString(190 * mm, y, f"{tax_sum:.2f} €")
    y -= 6 * mm
    c.setFont("Helvetica-Bold", 12)
    c.drawRightString(165 * mm, y, "Total:")
    c.drawRightString(190 * mm, y, f"{total_sum:.2f} €")

    c.showPage()
    c.save()
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes


# -----------------------------------------------------------------------------
# Routes
#
# These API endpoints allow you to programmatically interact with the
# application.  You can also build a simple front‑end on top of these by
# submitting forms via fetch/XHR.

# Create database tables once at startup.  The `before_first_request` decorator
# was removed in Flask 3【582101706213846†L169-L173】, so we explicitly initialize
# the database here using the application context.
with app.app_context():
    # Crear tablas base si no existen (solo desarrollo). En producción usar Alembic.
    app_env = (os.getenv('APP_ENV') or os.getenv('FLASK_ENV') or 'development').lower()
    allow_create_all = os.getenv('RUN_DB_CREATE_ALL', 'false').lower() in ('1', 'true', 'yes')
    allow_runtime_migrations = os.getenv('ALLOW_RUNTIME_MIGRATIONS', 'false').lower() in ('1', 'true', 'yes')

    if app_env != 'production' or allow_create_all:
        try:
            db.create_all()
        except Exception:
            # Evitar bloquear el arranque si hay desajustes temporales de esquema
            app.logger.warning('db.create_all() falló; confía en Alembic para crear/esquema')
            db.session.rollback()

    # Migración ligera (solo en desarrollo o si se permite explícitamente)
    if (app_env != 'production' and allow_runtime_migrations) or allow_create_all:
        try:
            insp = inspect(db.engine)
            cols = [c['name'] for c in insp.get_columns('product')]
            if 'is_active' not in cols:
                db.session.execute(text("ALTER TABLE product ADD COLUMN is_active BOOLEAN DEFAULT 1"))
                db.session.commit()
        except Exception:
            # No bloquear arranque si falla la migración ligera
            db.session.rollback()

    # Usuario admin inicial opcional (seguro en cualquier entorno si existen tablas)
    admin_user = os.getenv('ADMIN_USERNAME')
    admin_pass = os.getenv('ADMIN_PASSWORD')
    if admin_user and admin_pass:
        try:
            if not User.query.filter_by(username=admin_user).first():
                db.session.add(User(username=admin_user, password_hash=generate_password_hash(admin_pass)))
                db.session.commit()
        except Exception:
            db.session.rollback()


@app.route('/')
def index():
    """Home page displaying minimal UI for demonstration."""
    # The React front‑end embedded in index.html fetches clients via API, so
    # there's no need to pass any context here.  Just render the template.
    return render_template('index.html')


@app.route('/api/clients', methods=['POST'])
@jwt_required()
def create_client():
    """Create a new client via JSON request."""
    try:
        payload = ClientCreateRequest.model_validate(request.get_json(force=True))
    except PydValidationError as e:
        return jsonify({"error": e.errors()[0]['msg'] if e.errors() else 'invalid payload', "code": 400}), 400
    client = Client(
        name=payload.name,
        cif=payload.cif,
        address=payload.address,
        email=payload.email,
        phone=payload.phone,
        iban=payload.iban,
    )
    db.session.add(client)
    db.session.commit()
    return jsonify({'id': client.id}), 201


@app.route('/api/clients', methods=['GET'])
@jwt_required()
def list_clients():
    """Lista de clientes con paginación, búsqueda y orden.

    Parámetros opcionales:
      - limit (int): número máximo de elementos
      - offset (int): desplazamiento
      - q (str): término de búsqueda en nombre/cif/email/teléfono
      - sort (str): campo de ordenación (id, name, created_at, email, phone)
      - dir (str): dirección 'asc'|'desc' (por defecto 'desc')
    """
    limit = request.args.get('limit', type=int, default=10)
    offset = request.args.get('offset', type=int, default=0)
    q = (request.args.get('q') or '').strip()
    sort = (request.args.get('sort') or 'created_at').strip()
    direction = (request.args.get('dir') or 'desc').strip().lower()

    allowed_sort = {'id', 'name', 'created_at', 'email', 'phone'}
    if sort not in allowed_sort:
        sort = 'created_at'
    if direction not in {'asc', 'desc'}:
        direction = 'desc'

    query = Client.query
    if q:
        like = f"%{q}%"
        query = query.filter(
            db.or_(
                Client.name.ilike(like),
                Client.cif.ilike(like),
                Client.email.ilike(like),
                Client.phone.ilike(like),
            )
        )

    sort_col = getattr(Client, sort)
    if direction == 'desc':
        sort_col = sort_col.desc()
    query = query.order_by(sort_col)

    total = query.count()
    clients = query.offset(offset).limit(limit).all()

    items = []
    for c in clients:
        items.append({
            'id': c.id,
            'name': c.name,
            'cif': c.cif,
            'address': c.address,
            'email': c.email,
            'phone': c.phone,
            'iban': c.iban,
            'created_at': c.created_at.isoformat() if getattr(c, 'created_at', None) else None,
        })
    return jsonify({'items': items, 'total': total})


@app.route('/api/invoices', methods=['POST'])
@jwt_required()
def create_invoice():
    """Create a new invoice or proforma and its items.

    Número se asigna automáticamente según el tipo.
    """
    try:
        payload = InvoiceCreateRequest.model_validate(request.get_json(force=True))
    except PydValidationError as e:
        return jsonify({"error": e.errors()[0]['msg'] if e.errors() else 'invalid payload', "code": 400}), 400
    date_str = payload.date
    invoice_type = payload.type
    client_id = payload.client_id
    notes = payload.notes or ''
    payment_method = payload.payment_method
    items_data = [i.model_dump() for i in payload.items]
    # Convert date string to date object
    date_obj = datetime.strptime(date_str, '%Y-%m-%d').date()
    # Normalize payment method
    allowed_pm = {'efectivo','bizum','transferencia'}
    if invoice_type != 'factura':
        payment_method = None
    else:
        pm = (payment_method or '').strip().lower()
        payment_method = pm if pm in allowed_pm else 'efectivo'
    # Compute totals
    subtotal, tax_amount, total = calculate_totals(items_data)
    # Asignar número automáticamente según fecha indicada (reinicia por año/mes)
    number = _next_sequence_atomic(invoice_type, datetime.strptime(date_str, '%Y-%m-%d'))
    paid_flag = bool(payload.paid)
    if invoice_type != 'factura':
        paid_flag = False
    invoice = Invoice(
        number=number,
        date=date_obj,
        type=invoice_type,
        client_id=client_id,
        notes=notes,
        payment_method=payment_method,
        total=total,
        tax_total=tax_amount,
        paid=paid_flag
    )
    db.session.add(invoice)
    try:
        db.session.flush()  # flush to obtain invoice.id
    except IntegrityError:
        db.session.rollback()
        return jsonify({'error': 'El número de factura ya existe'}), 409
    # Validar stock de productos (solo para facturas reales)
    products_to_decrement = []  # (product, qty)
    if invoice_type == 'factura':
        for item in items_data:
            pid = item.get('product_id')
            if pid:
                prod = Product.query.get(pid)
                if not prod:
                    return jsonify({'error': f'Producto {pid} no existe', 'code': 400}), 400
                qty = int(item['units'])
                if (prod.stock_qty or 0) < qty:
                    return jsonify({'error': f'Sin stock suficiente para producto {pid}', 'code': 409}), 409
                products_to_decrement.append((prod, qty))

    for item in items_data:
        line = InvoiceItem(
            invoice_id=invoice.id,
            product_id=item.get('product_id'),
            description=item['description'],
            units=item['units'],
            unit_price=item['unit_price'],
            tax_rate=item['tax_rate'],
            subtotal=item['units'] * item['unit_price'],
            total=item['units'] * item['unit_price'] * (1 + item['tax_rate'] / 100)
        )
        db.session.add(line)
    # Descontar stock y registrar movimiento (solo 'factura')
    if invoice_type == 'factura':
        for prod, qty in products_to_decrement:
            prod.stock_qty = int(prod.stock_qty or 0) - int(qty)
            db.session.add(StockMovement(product_id=prod.id, qty=-int(qty), type='sale', invoice_id=invoice.id))
    db.session.commit()
    return jsonify({
        'id': invoice.id,
        'number': invoice.number,
        'date': invoice.date.isoformat(),
        'type': invoice.type,
        'client_id': invoice.client_id,
        'notes': invoice.notes,
        'payment_method': invoice.payment_method,
        'total': invoice.total,
        'tax_total': invoice.tax_total,
        'paid': bool(invoice.paid),
        'items': [
            {
                'description': it.description,
                'units': it.units,
                'unit_price': it.unit_price,
                'tax_rate': it.tax_rate,
                'product_id': it.product_id,
                'subtotal': it.subtotal,
                'total': it.total,
            }
            for it in invoice.items
        ],
    }), 201


@app.route('/api/invoices', methods=['GET'])
@jwt_required()
def list_invoices():
    """Listado de facturas/proformas con filtros, paginación, búsqueda y orden.

    Parámetros opcionales:
      - month, year (int): filtrar por mes/año
      - limit (int), offset (int)
      - q (str): búsqueda por número o nombre de cliente
      - sort (str): id, date, total, tax_total, number
      - dir (str): asc|desc (por defecto desc)
    """
    month = request.args.get('month')
    year = request.args.get('year')
    limit = request.args.get('limit', type=int, default=10)
    offset = request.args.get('offset', type=int, default=0)
    q = (request.args.get('q') or '').strip()
    sort = (request.args.get('sort') or 'date').strip()
    direction = (request.args.get('dir') or 'desc').strip().lower()

    allowed_sort = {'id', 'date', 'total', 'tax_total', 'number'}
    if sort not in allowed_sort:
        sort = 'date'
    if direction not in {'asc', 'desc'}:
        direction = 'desc'

    query = Invoice.query

    # Filtro por mes/año
    if month and year:
        try:
            m = int(month)
            y = int(year)
            query = query.filter(db.extract('month', Invoice.date) == m)
            query = query.filter(db.extract('year', Invoice.date) == y)
        except ValueError:
            return jsonify({'error': 'Month and year must be integers'}), 400

    # Búsqueda por número o nombre de cliente
    if q:
        like = f"%{q}%"
        # Join para poder buscar por nombre de cliente
        query = query.join(Client, Client.id == Invoice.client_id)
        query = query.filter(db.or_(Invoice.number.ilike(like), Client.name.ilike(like)))

    # Orden
    sort_col = getattr(Invoice, sort)
    if direction == 'desc':
        sort_col = sort_col.desc()
    query = query.order_by(sort_col)

    total = query.count()
    invoices = query.offset(offset).limit(limit).all()

    items = []
    for inv in invoices:
        items.append({
            'id': inv.id,
            'number': inv.number,
            'date': inv.date.isoformat(),
            'client_id': inv.client_id,
            'type': inv.type,
            'payment_method': inv.payment_method,
            'total': inv.total,
            'tax_total': inv.tax_total,
            'paid': bool(inv.paid)
        })
    return jsonify({'items': items, 'total': total})


@app.route('/api/invoices/next_number')
@jwt_required()
def get_next_number():
    """Devuelve el siguiente número formateado para un tipo dado (factura|proforma)."""
    doc_type = request.args.get('type', 'factura')
    if doc_type not in ('factura', 'proforma'):
        return jsonify({'error': 'type inválido'}), 400
    # Permite opcionalmente fecha base para preview (YYYY-MM-DD); por defecto hoy
    date_str = request.args.get('date')
    base = datetime.strptime(date_str, '%Y-%m-%d') if date_str else datetime.utcnow()
    seq = DocumentSequence.query.filter_by(doc_type=doc_type, year=base.year, month=base.month).first()
    last = seq.last_number if seq else 0
    next_formatted = _format_number_for_type(doc_type, last + 1, base.year, base.month)
    return jsonify({'next_number': next_formatted})


@app.route('/api/invoices/<int:invoice_id>', methods=['GET'])
@jwt_required()
def get_invoice(invoice_id):
    inv = Invoice.query.get_or_404(invoice_id)
    return jsonify({
        'id': inv.id,
        'number': inv.number,
        'date': inv.date.isoformat(),
        'client_id': inv.client_id,
        'type': inv.type,
        'notes': inv.notes,
        'payment_method': inv.payment_method,
        'total': inv.total,
        'tax_total': inv.tax_total,
        'paid': bool(inv.paid),
        'items': [
            {
                'description': it.description,
                'units': it.units,
                'unit_price': it.unit_price,
                'tax_rate': it.tax_rate,
                'subtotal': it.subtotal,
                'total': it.total,
            }
            for it in inv.items
        ],
    })


@app.route('/api/invoices/<int:invoice_id>', methods=['DELETE'])
@jwt_required()
def delete_invoice(invoice_id):
    inv = Invoice.query.get_or_404(invoice_id)
    # Política A: impedir borrar si hay ventas de productos asociadas
    # Bloquea si la factura es real y tiene líneas con product_id o movimientos registrados
    if inv.type == 'factura':
        has_product_lines = any((it.product_id is not None) for it in inv.items)
        has_stock_moves = StockMovement.query.filter_by(invoice_id=inv.id).count() > 0
        if has_product_lines or has_stock_moves:
            return jsonify({'error': 'No se puede eliminar: factura con productos vendidos'}), 409
    for it in list(inv.items):
        db.session.delete(it)
    db.session.delete(inv)
    db.session.commit()
    return jsonify({'status': 'deleted'})


@app.route('/api/clients/<int:client_id>', methods=['PUT'])
@jwt_required()
def update_client(client_id):
    client = Client.query.get_or_404(client_id)
    data = request.get_json(force=True)
    for field in ['name', 'cif', 'address', 'email', 'phone', 'iban']:
        if field in data:
            setattr(client, field, data[field])
    db.session.commit()
    return jsonify({'status': 'ok'})


@app.route('/api/clients/<int:client_id>', methods=['DELETE'])
@jwt_required()
def delete_client(client_id):
    """Elimina un cliente si no tiene facturas asociadas."""
    client = Client.query.get_or_404(client_id)
    # Evitar borrar clientes con facturas relacionadas para no romper integridad
    has_invoices = Invoice.query.filter_by(client_id=client.id).count() > 0
    if has_invoices:
        return jsonify({'error': 'No se puede eliminar: el cliente tiene facturas asociadas'}), 409
    db.session.delete(client)
    db.session.commit()
    return jsonify({'status': 'deleted'})


def _client_upload_dir(client_id: int) -> str:
    base = os.path.join(UPLOADS_ROOT, str(client_id))
    os.makedirs(os.path.join(base, 'documents'), exist_ok=True)
    os.makedirs(os.path.join(base, 'images'), exist_ok=True)
    return base


def _allowed_category_and_mime(filename: str, content_type: str) -> tuple[str, bool]:
    name = filename.lower()
    if name.endswith('.pdf') and content_type in ('application/pdf', 'application/octet-stream'):
        return 'document', True
    if any(name.endswith(ext) for ext in ('.jpg', '.jpeg', '.png', '.webp')) and content_type.startswith('image/'):
        return 'image', True
    return 'other', False


@app.route('/api/clients/<int:client_id>/documents', methods=['GET'])
@jwt_required()
def list_client_documents(client_id):
    Client.query.get_or_404(client_id)
    docs = (ClientDocument.query
            .filter_by(client_id=client_id)
            .order_by(ClientDocument.uploaded_at.desc())
            .all())
    return jsonify([
        {
            'id': d.id,
            'category': d.category,
            'filename': d.filename,
            'content_type': d.content_type,
            'size_bytes': d.size_bytes,
            'uploaded_at': d.uploaded_at.isoformat(),
        } for d in docs
    ])


@app.route('/api/clients/<int:client_id>/documents', methods=['POST'])
@jwt_required()
@limiter.limit("20 per minute")
def upload_client_document(client_id):
    Client.query.get_or_404(client_id)
    if 'file' not in request.files:
        return jsonify({'error': 'Fichero requerido (campo file)'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nombre de fichero vacío'}), 400
    safe_name = secure_filename(file.filename)
    category, ok = _allowed_category_and_mime(safe_name, file.mimetype or '')
    if not ok:
        return jsonify({'error': 'Tipo de archivo no permitido (PDF o imagen)'}), 400
    # Guardar con nombre único para evitar colisiones
    base_dir = _client_upload_dir(client_id)
    subdir = 'documents' if category == 'document' else 'images'
    unique_name = f"{uuid4().hex}_{safe_name}"
    stored_rel = os.path.join(str(client_id), subdir, unique_name)
    stored_abs = os.path.join(UPLOADS_ROOT, stored_rel)
    file.save(stored_abs)
    size_bytes = os.path.getsize(stored_abs)
    doc = ClientDocument(
        client_id=client_id,
        category=category,
        filename=safe_name,
        stored_path=stored_rel,
        content_type=file.mimetype or 'application/octet-stream',
        size_bytes=size_bytes,
    )
    db.session.add(doc)
    db.session.commit()
    return jsonify({'id': doc.id}), 201


@app.route('/api/clients/<int:client_id>/documents/<int:doc_id>', methods=['GET'])
@jwt_required()
@limiter.limit("100 per minute")
def get_client_document(client_id, doc_id):
    doc = ClientDocument.query.filter_by(id=doc_id, client_id=client_id).first_or_404()
    abs_path = os.path.join(UPLOADS_ROOT, doc.stored_path)
    if not os.path.isfile(abs_path):
        abort(404)
    dl_flag = (request.args.get('dl', '') or '').lower() in ('1','true','yes')
    resp = send_file(
        abs_path,
        mimetype=doc.content_type,
        as_attachment=dl_flag,
        download_name=doc.filename,
    )
    resp.headers['Cache-Control'] = 'private, no-store'
    return resp


@app.route('/api/clients/<int:client_id>/documents/<int:doc_id>', methods=['DELETE'])
@jwt_required()
def delete_client_document(client_id, doc_id):
    doc = ClientDocument.query.filter_by(id=doc_id, client_id=client_id).first_or_404()
    abs_path = os.path.join(UPLOADS_ROOT, doc.stored_path)
    try:
        if os.path.isfile(abs_path):
            os.remove(abs_path)
    except Exception:
        pass
    db.session.delete(doc)
    db.session.commit()
    return jsonify({'status': 'deleted'})


@app.route('/api/clients/<int:client_id>/invoices', methods=['GET'])
@jwt_required()
def list_invoices_by_client(client_id):
    Client.query.get_or_404(client_id)
    limit = request.args.get('limit', type=int, default=10)
    offset = request.args.get('offset', type=int, default=0)
    sort = (request.args.get('sort') or 'id').strip()
    direction = (request.args.get('dir') or 'desc').strip().lower()
    allowed_sort = {'id', 'date', 'total', 'tax_total', 'number'}
    if sort not in allowed_sort:
        sort = 'id'
    if direction not in {'asc', 'desc'}:
        direction = 'desc'
    sort_col = getattr(Invoice, sort)
    if direction == 'desc':
        sort_col = sort_col.desc()
    q = Invoice.query.filter_by(client_id=client_id).order_by(sort_col)
    total = q.count()
    invs = q.offset(offset).limit(limit).all()
    return jsonify({
        'items': [
            {
                'id': inv.id,
                'number': inv.number,
                'date': inv.date.isoformat(),
                'type': inv.type,
                'total': inv.total,
                'tax_total': inv.tax_total,
            } for inv in invs
        ],
        'total': total,
    })


@app.get('/api/contracts/templates')
@jwt_required()
def list_contract_templates():
    """List available contract templates."""
    templates = [
        {
            'id': 'compraventa',
            'name': 'Contrato de Compraventa',
            'filename': 'Contrato_Compraventa_Plazos_NIOXTEC_v5.docx',
            'description': 'Contrato de compraventa con plazos'
        },
        {
            'id': 'renting',
            'name': 'Contrato de Renting',
            'filename': 'Plantilla_Contrato_Renting_Firma_Datos_v2.docx',
            'description': 'Contrato de renting con firma y datos'
        }
    ]
    return jsonify(templates)


# -----------------------------
# Products API
# -----------------------------

def _product_to_dict(p: 'Product') -> dict:
    return {
        'id': p.id,
        'category': p.category,
        'model': p.model,
        'sku': p.sku,
        'stock_qty': p.stock_qty,
        'price_net': p.price_net,
        'tax_rate': p.tax_rate,
        'features': p.features or {},
        'images': p.images or [],  # Incluir imágenes
        'created_at': p.created_at.isoformat() if p.created_at else None,
    'is_active': bool(getattr(p, 'is_active', True)),
    }


@app.route('/api/products', methods=['POST'])
@jwt_required()
def create_product():
    data = request.get_json(force=True)
    for field in ['category', 'model']:
        if not data.get(field):
            return jsonify({'error': f'{field} requerido'}), 400
    try:
        stock = int(data.get('stock_qty') or 0)
        price_net = float(data.get('price_net') or 0)
        tax_rate = float(data.get('tax_rate') or 21.0)
    except Exception:
        return jsonify({'error': 'stock_qty/price_net/tax_rate inválidos'}), 400
    p = Product(
        category=str(data['category']).strip(),
        model=str(data['model']).strip(),
        sku=(data.get('sku') or None),
        stock_qty=max(0, stock),
        price_net=max(0.0, price_net),
        tax_rate=max(0.0, min(100.0, tax_rate)),
        features=data.get('features') or {},
    )
    db.session.add(p)
    db.session.commit()
    return jsonify({'id': p.id}), 201


@app.route('/api/products', methods=['GET', 'OPTIONS'])
@jwt_required(optional=True)
def list_products():
    limit = request.args.get('limit', type=int, default=10)
    offset = request.args.get('offset', type=int, default=0)
    q = (request.args.get('q') or '').strip()
    category = (request.args.get('category') or '').strip()
    model = (request.args.get('model') or '').strip()
    sort = (request.args.get('sort') or 'created_at').strip()
    direction = (request.args.get('dir') or 'desc').strip().lower()
    active_param = (request.args.get('active') or '').strip()
    allowed_sort = {'id','category','model','sku','stock_qty','price_net','created_at'}
    if sort not in allowed_sort: sort = 'created_at'
    if direction not in {'asc','desc'}: direction = 'desc'
    query = Product.query
    if category:
        query = query.filter(Product.category.ilike(category))
    if model:
        query = query.filter(Product.model.ilike(model))
    if q:
        like = f"%{q}%"
        query = query.filter(db.or_(Product.model.ilike(like), Product.category.ilike(like), Product.sku.ilike(like)))
    # Por defecto devolver solo activos; si active=0 devuelve archivados; si active=1 activos
    if active_param in {'0','1'}:
        query = query.filter(Product.is_active == (active_param == '1'))
    else:
        query = query.filter(Product.is_active == True)  # noqa: E712
    sort_col = getattr(Product, sort)
    if direction == 'desc': sort_col = sort_col.desc()
    query = query.order_by(sort_col)
    total = query.count()
    rows = query.offset(offset).limit(limit).all()
    return jsonify({'items': [_product_to_dict(p) for p in rows], 'total': total})


@app.route('/api/products/<int:pid>', methods=['GET'])
@jwt_required()
def get_product(pid):
    p = Product.query.get_or_404(pid)
    return jsonify(_product_to_dict(p))


@app.route('/api/products/<int:pid>', methods=['PUT'])
@jwt_required()
def update_product(pid):
    p = Product.query.get_or_404(pid)
    data = request.get_json(force=True)
    # Update simple string fields, but treat SKU specially: empty -> NULL
    for field in ['category', 'model', 'features', 'images']:
        if field in data:
            setattr(p, field, data[field])
            # Si es images, marcar como modificado para SQLAlchemy
            if field == 'images':
                from sqlalchemy.orm.attributes import flag_modified
                flag_modified(p, 'images')
    if 'sku' in data:
        # store NULL in DB when client sends empty/blank sku to avoid
        # violating UNIQUE constraint for empty strings
        sku = data.get('sku')
        if isinstance(sku, str):
            sku = sku.strip() or None
        else:
            sku = sku or None
        p.sku = sku
    if 'stock_qty' in data:
        try:
            p.stock_qty = max(0, int(data['stock_qty']))
        except Exception:
            return jsonify({'error': 'stock_qty inválido'}), 400
    if 'price_net' in data:
        try:
            p.price_net = max(0.0, float(data['price_net']))
        except Exception:
            return jsonify({'error': 'price_net inválido'}), 400
    if 'tax_rate' in data:
        try:
            tr = float(data['tax_rate'])
            if tr < 0 or tr > 100: raise ValueError()
            p.tax_rate = tr
        except Exception:
            return jsonify({'error': 'tax_rate inválido (0–100)'}), 400
    if 'is_active' in data:
        try:
            p.is_active = bool(data['is_active'])
        except Exception:
            return jsonify({'error': 'is_active inválido'}), 400
    try:
        db.session.commit()
    except IntegrityError as e:
        # Likely a UNIQUE constraint on sku or similar
        db.session.rollback()
        return jsonify({'error': 'Integrity error updating product (possible duplicate SKU)'}), 409
    return jsonify({'status': 'ok'})


@app.route('/api/products/<int:pid>', methods=['DELETE'])
@jwt_required()
def delete_product(pid):
    p = Product.query.get_or_404(pid)
    # Impedir borrar si tiene ventas o está referenciado en invoice items
    ref_count = InvoiceItem.query.filter_by(product_id=p.id).count()
    mov_count = StockMovement.query.filter_by(product_id=p.id).count()
    if ref_count > 0 or mov_count > 0:
        return jsonify({'error': 'No se puede borrar: referenciado en ventas o con movimientos'}), 409
    db.session.delete(p)
    db.session.commit()
    return jsonify({'status': 'deleted'})


@app.route('/api/products/summary', methods=['GET', 'OPTIONS'])
@jwt_required(optional=True)
def products_summary():
    # Agregados por categoría y por modelo, filtrando por activos/archivados
    active_param = (request.args.get('active') or '1').strip()
    where_clause = "WHERE is_active = :active" if active_param in {'0','1'} else "WHERE is_active = 1"
    cats = db.session.execute(text(
        f"SELECT category, COUNT(*) as total FROM product {where_clause} GROUP BY category ORDER BY category"
    ), {'active': 1 if active_param != '0' else 0}).fetchall()
    by_cat = {}
    for c, t in cats:
        # Obtener suma de stock_qty y cantidad de productos por modelo
        rows = db.session.execute(text(
            f"""SELECT model, 
                       COUNT(*) as cnt, 
                       SUM(COALESCE(stock_qty, 0)) as stock_total 
                FROM product {where_clause} AND category = :c 
                GROUP BY model 
                ORDER BY model"""
        ), {'c': c, 'active': 1 if active_param != '0' else 0}).fetchall()
        by_cat[c] = {
            'category': c,
            'total': int(t or 0),
            'models': [{'model': m, 'count': int(cnt or 0), 'stock_total': int(stock_total or 0)} for m, cnt, stock_total in rows]
        }
    return jsonify({'categories': list(by_cat.values())})


@app.route('/api/products/<int:pid>/adjust_stock', methods=['POST'])
@jwt_required()
def adjust_product_stock(pid: int):
    """Ajuste manual del stock de un producto.

    Payload JSON:
      - qty: int (positivo para entrada, negativo para salida)
      - type: 'manual'|'adjust' (opcional, por defecto 'adjust')
    Reglas:
      - No permite dejar el stock en negativo.
      - Registra un StockMovement con la cantidad indicada.
    """
    p = Product.query.get_or_404(pid)
    # No permitir ajustes si está archivado
    if not getattr(p, 'is_active', True):
        return jsonify({'error': 'Producto archivado: no se puede ajustar stock'}), 409
    data = request.get_json(force=True)
    try:
        qty = int(data.get('qty'))
    except Exception:
        return jsonify({'error': 'qty requerido (entero no nulo)'}), 400
    if qty == 0:
        return jsonify({'error': 'qty no puede ser 0'}), 400
    mv_type = (data.get('type') or 'adjust').strip().lower()
    if mv_type not in {'manual', 'adjust'}:
        mv_type = 'adjust'

    new_qty = int(p.stock_qty or 0) + qty
    if new_qty < 0:
        return jsonify({'error': 'stock resultante no puede ser negativo'}), 409

    p.stock_qty = new_qty
    db.session.add(StockMovement(product_id=p.id, qty=int(qty), type=mv_type))
    db.session.commit()
    return jsonify({'status': 'ok', 'stock_qty': p.stock_qty})

@app.route('/api/products/upload-image', methods=['POST', 'OPTIONS'])
@jwt_required(optional=True)
def upload_product_image():
    """Subir una imagen para un producto."""
    # Manejar preflight CORS
    if request.method == 'OPTIONS':
        response = jsonify({'status': 'ok'})
        response.headers.add('Access-Control-Allow-Origin', request.headers.get('Origin', '*'))
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
        response.headers.add('Access-Control-Allow-Methods', 'POST,OPTIONS')
        return response, 200
    
    # Verificar JWT para POST
    verify_jwt_in_request()
    
    import os
    from werkzeug.utils import secure_filename
    
    print("📸 DEBUG: Iniciando subida de imagen...")
    print(f"📸 DEBUG: request.files keys: {list(request.files.keys())}")
    print(f"📸 DEBUG: request.form keys: {list(request.form.keys())}")
    
    # Verificar que se envió un archivo
    if 'file' not in request.files:
        print("❌ DEBUG: No se encontró 'file' en request.files")
        return jsonify({'error': 'No se envió ningún archivo'}), 400
    
    file = request.files['file']
    product_id = request.form.get('product_id')
    
    print(f"📸 DEBUG: file.filename = {file.filename}")
    print(f"📸 DEBUG: product_id = {product_id}")
    
    if file.filename == '':
        return jsonify({'error': 'Nombre de archivo vacío'}), 400
    
    if not product_id:
        return jsonify({'error': 'product_id es requerido'}), 400
    
    # Verificar que el producto existe
    product = Product.query.get_or_404(int(product_id))
    
    # Validar extensión de archivo
    allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
    ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
    if ext not in allowed_extensions:
        return jsonify({'error': f'Extensión no permitida. Usa: {", ".join(allowed_extensions)}'}), 400
    
    # Crear directorio si no existe
    upload_dir = os.path.join('static', 'uploads', 'products')
    os.makedirs(upload_dir, exist_ok=True)
    
    # Generar nombre único para el archivo
    filename = secure_filename(file.filename)
    unique_filename = f"{product_id}_{int(time.time())}_{filename}"
    file_path = os.path.join(upload_dir, unique_filename)
    
    print(f"📸 DEBUG: Guardando en: {file_path}")
    
    # Guardar archivo
    file.save(file_path)
    
    # URL relativa para acceder a la imagen
    image_url = f"/static/uploads/products/{unique_filename}"
    
    print(f"✅ DEBUG: Imagen guardada exitosamente: {image_url}")
    
    # Guardar en la base de datos (actualizar campo images del producto)
    # Verificar si el producto tiene el atributo images y si es None, inicializarlo como lista
    try:
        current_images = product.images if hasattr(product, 'images') and product.images else []
    except Exception as e:
        print(f"⚠️ DEBUG: Error al leer product.images: {e}")
        current_images = []
    
    current_images.append({
        'url': image_url,
        'alt': file.filename
    })
    product.images = current_images
    
    # Marcar el atributo como modificado para que SQLAlchemy lo actualice
    from sqlalchemy.orm.attributes import flag_modified
    flag_modified(product, 'images')
    
    db.session.commit()
    
    print(f"✅ DEBUG: Base de datos actualizada")
    
    return jsonify({
        'url': image_url,
        'alt': file.filename,
        'status': 'ok'
    })

@app.get('/api/contracts/templates/<template_id>/placeholders')
@jwt_required()
def get_template_placeholders(template_id):
    """Get placeholders from a specific template."""
    templates = {
        'compraventa': 'Contrato_Compraventa_Plazos_NIOXTEC_v5.docx',
        'renting': 'Plantilla_Contrato_Renting_Firma_Datos_v2.docx'
    }
    
    if template_id not in templates:
        return jsonify({'error': 'Template not found'}), 404
    
    result = extract_placeholders_from_docx(templates[template_id])
    
    if 'error' in result:
        return jsonify(result), 400
    
    return jsonify(result)

@app.get('/api/company/config')
@jwt_required()
def get_company_config():
    """Get company configuration for contract generation."""
    company = CompanyConfig.query.first()
    if not company:
        # Fallback a variables de entorno si no existe fila en DB
        env_company = _company_from_env()
        return jsonify({
            'name': env_company.name,
            'cif': env_company.cif,
            'address': env_company.address,
            'city': env_company.city or '',
            'province': env_company.province or '',
            'email': env_company.email,
            'phone': env_company.phone,
            'iban': env_company.iban or '',
            'website': env_company.website or ''
        })
    
    return jsonify({
        'name': company.name,
        'cif': company.cif,
        'address': company.address,
        'city': company.city or '',
        'province': company.province or '',
        'email': company.email,
        'phone': company.phone,
        'iban': company.iban or '',
        'website': company.website or ''
    })

@app.route('/api/invoices/<int:invoice_id>', methods=['PUT'])
@jwt_required()
def update_invoice(invoice_id):
    inv = Invoice.query.get_or_404(invoice_id)
    data = request.get_json(force=True)
    # Basic fields
    # Ya no permitimos cambiar el número arbitrariamente para mantener la secuencia
    if 'date' in data:
        inv.date = datetime.strptime(data['date'], '%Y-%m-%d').date()
    if 'type' in data:
        inv.type = data['type']
    if 'client_id' in data:
        inv.client_id = int(data['client_id'])
    if 'payment_method' in data:
        pm = (data.get('payment_method') or '').strip().lower()
        inv.payment_method = pm if inv.type == 'factura' and pm in {'efectivo','bizum','transferencia'} else (None if inv.type!='factura' else 'efectivo')
    inv.notes = data.get('notes', inv.notes)
    if 'paid' in data:
        inv.paid = bool(data.get('paid')) if inv.type == 'factura' else False
    # Reglas de edición de líneas
    items_data = data.get('items')
    # Política: si es factura y hay líneas vinculadas a productos, no permitir edición de líneas
    if items_data is not None and inv.type == 'factura':
        # Detectar si la factura actual tiene líneas con product_id
        if any((it.product_id is not None) for it in inv.items):
            return jsonify({'error': 'No se pueden editar las líneas: factura con productos vinculados'}), 409
        # Además, impedir que el payload intente establecer product_id vía edición
        try:
            if any(('product_id' in (item or {})) for item in items_data):
                return jsonify({'error': 'No se pueden vincular productos al editar una factura existente'}), 409
        except Exception:
            pass
    if items_data is not None:
        for it in list(inv.items):
            db.session.delete(it)
        subtotal, tax_amount, total = calculate_totals(items_data)
        inv.total = total
        inv.tax_total = tax_amount
        db.session.flush()
        for item in items_data:
            line = InvoiceItem(
                invoice_id=inv.id,
                description=item['description'],
                units=item['units'],
                unit_price=item['unit_price'],
                tax_rate=item['tax_rate'],
                subtotal=item['units'] * item['unit_price'],
                total=item['units'] * item['unit_price'] * (1 + item['tax_rate'] / 100)
            )
            db.session.add(line)
    try:
        db.session.commit()
    except IntegrityError:
        db.session.rollback()
        return jsonify({'error': 'El número de factura ya existe'}), 409
    return jsonify({'status': 'ok'})


@app.route('/api/invoices/<int:invoice_id>/paid', methods=['PATCH'])
@jwt_required()
def update_invoice_paid(invoice_id):
    inv = Invoice.query.get_or_404(invoice_id)
    if inv.type != 'factura':
        return jsonify({'error': 'Solo se puede marcar pagado en facturas'}), 400
    data = request.get_json(silent=True) or {}
    paid_value = bool(data.get('paid'))
    inv.paid = paid_value
    db.session.commit()
    return jsonify({'status': 'ok', 'paid': bool(inv.paid)})


@app.route('/api/invoices/<int:invoice_id>/convert', methods=['PATCH'])
@jwt_required()
def convert_proforma(invoice_id):
    """Crea una nueva factura a partir de una proforma y conserva la proforma.

    - Asigna número secuencial nuevo a la factura
    - Copia líneas/items
    - La proforma original NO se modifica ni elimina
    """
    inv = Invoice.query.get_or_404(invoice_id)
    if inv.type == 'factura':
        return jsonify({'error': 'El documento ya es una factura'}), 400

    # Payload opcional (para método de pago)
    data = request.get_json(silent=True) or {}
    allowed_pm = {'efectivo', 'bizum', 'transferencia'}
    pm = (data.get('payment_method') or 'efectivo').strip().lower()
    payment_method = pm if pm in allowed_pm else 'efectivo'

    # 1) Validar stock por cada línea con product_id
    items_with_product = [it for it in inv.items if getattr(it, 'product_id', None)]
    products_and_qty: list[tuple[Product, int]] = []
    for it in items_with_product:
        prod = Product.query.get(it.product_id)
        if not prod:
            return jsonify({'error': f'Producto {it.product_id} no existe', 'code': 400}), 400
        qty = int(it.units)
        if int(prod.stock_qty or 0) < qty:
            return jsonify({'error': f'Sin stock suficiente para producto {prod.id}', 'code': 409}), 409
        products_and_qty.append((prod, qty))

    # 2) Crear nueva factura (mantener proforma original)
    number = _next_sequence_atomic('factura')
    new_inv = Invoice(
        number=number,
        date=datetime.utcnow().date(),
        type='factura',
        client_id=inv.client_id,
        notes=inv.notes,
        payment_method=payment_method,
        total=inv.total,
        tax_total=inv.tax_total,
        paid=False,
    )
    db.session.add(new_inv)
    db.session.flush()

    # 3) Copiar líneas de la proforma a la nueva factura
    for it in inv.items:
        line_subtotal = it.units * it.unit_price
        line_total = line_subtotal * (1 + it.tax_rate / 100)
        db.session.add(InvoiceItem(
            invoice_id=new_inv.id,
            product_id=getattr(it, 'product_id', None),
            description=it.description,
            units=it.units,
            unit_price=it.unit_price,
            tax_rate=it.tax_rate,
            subtotal=line_subtotal,
            total=line_total,
        ))

    # 4) Descontar stock y registrar movimientos asociados a la nueva factura
    for prod, qty in products_and_qty:
        prod.stock_qty = int(prod.stock_qty or 0) - int(qty)
        db.session.add(StockMovement(product_id=prod.id, qty=-int(qty), type='sale', invoice_id=new_inv.id))
    try:
        db.session.commit()
    except IntegrityError:
        db.session.rollback()
        return jsonify({'error': 'El número de factura ya existe'}), 409

    return jsonify({
        'status': 'ok',
        'number': new_inv.number,
        'invoice': {
            'id': new_inv.id,
            'number': new_inv.number,
            'date': new_inv.date.isoformat(),
            'client_id': new_inv.client_id,
            'type': new_inv.type,
            'payment_method': new_inv.payment_method,
            'total': new_inv.total,
            'tax_total': new_inv.tax_total,
            'paid': bool(new_inv.paid),
        }
    })


@app.route('/api/reports/summary')
@jwt_required()
def reports_summary():
    year = request.args.get('year', type=int, default=datetime.utcnow().year)
    # Sum totals by month for invoices type 'factura'
    rows = []
    try:
        rows = (
            db.session.query(db.extract('month', Invoice.date).label('month'), db.func.sum(Invoice.total))
            .filter(db.extract('year', Invoice.date) == year)
            .filter(Invoice.type == 'factura')
            .filter(Invoice.paid.is_(True))
            .group_by('month')
            .order_by('month')
            .all()
        )
    except Exception:
        # SQLite fallback using strftime
        rows = db.session.execute(
            text("""
                SELECT CAST(STRFTIME('%m', date) AS INTEGER) AS month, SUM(total)
                FROM invoice
                WHERE type = 'factura' AND paid = 1 AND CAST(STRFTIME('%Y', date) AS INTEGER) = :year
                GROUP BY month
                ORDER BY month
            """), { 'year': year }
        ).fetchall()
    by_month = {int(m): float(t or 0) for m, t in rows}
    total_year = sum(by_month.values())
    return jsonify({'year': year, 'by_month': by_month, 'total_year': total_year})


@app.route('/api/reports/heatmap')
@jwt_required()
def reports_heatmap():
    year = request.args.get('year', type=int)
    month = request.args.get('month', type=int)
    if not year or not month:
        return jsonify({'error': 'Parámetros year y month requeridos'}), 400
    try:
        rows = (
            db.session.query(Invoice.date, db.func.sum(Invoice.total))
            .filter(db.extract('year', Invoice.date) == year)
            .filter(db.extract('month', Invoice.date) == month)
            .filter(Invoice.type == 'factura')
            .filter(Invoice.paid.is_(True))
            .group_by(Invoice.date)
            .all()
        )
    except Exception:
        rows = db.session.execute(
            text("""
                SELECT date as d, SUM(total) as t
                FROM invoice
                WHERE type = 'factura'
                  AND paid = 1
                  AND CAST(STRFTIME('%Y', date) AS INTEGER) = :year
                  AND CAST(STRFTIME('%m', date) AS INTEGER) = :month
                GROUP BY d
                ORDER BY d
            """), { 'year': year, 'month': month }
        ).fetchall()
    by_day = {d.strftime('%Y-%m-%d'): float(t or 0) for d, t in rows}
    return jsonify({'year': year, 'month': month, 'by_day': by_day})


@app.route('/api/reports/expenses_summary')
@jwt_required()
def reports_expenses_summary():
    year = request.args.get('year', type=int, default=datetime.utcnow().year)
    # Sum totals by month for expenses
    rows = []
    try:
        rows = (
            db.session.query(db.extract('month', Expense.date).label('month'), db.func.sum(Expense.total))
            .filter(db.extract('year', Expense.date) == year)
            .group_by('month')
            .order_by('month')
            .all()
        )
    except Exception:
        # SQLite fallback using strftime
        rows = db.session.execute(
            text("""
                SELECT CAST(STRFTIME('%m', date) AS INTEGER) AS month, SUM(total)
                FROM expense
                WHERE CAST(STRFTIME('%Y', date) AS INTEGER) = :year
                GROUP BY month
                ORDER BY month
            """), { 'year': year }
        ).fetchall()
    by_month = {int(m): float(t or 0) for m, t in rows}
    total_year = sum(by_month.values())
    return jsonify({'year': year, 'by_month': by_month, 'total_year': total_year})


@app.route('/api/reports/expenses_heatmap')
@jwt_required()
def reports_expenses_heatmap():
    year = request.args.get('year', type=int)
    month = request.args.get('month', type=int)
    if not year or not month:
        return jsonify({'error': 'Parámetros year y month requeridos'}), 400
    try:
        rows = (
            db.session.query(Expense.date, db.func.sum(Expense.total))
            .filter(db.extract('year', Expense.date) == year)
            .filter(db.extract('month', Expense.date) == month)
            .group_by(Expense.date)
            .all()
        )
    except Exception:
        rows = db.session.execute(
            text("""
                SELECT date as d, SUM(total) as t
                FROM expense
                WHERE CAST(STRFTIME('%Y', date) AS INTEGER) = :year
                  AND CAST(STRFTIME('%m', date) AS INTEGER) = :month
                GROUP BY d
                ORDER BY d
            """), { 'year': year, 'month': month }
        ).fetchall()
    by_day = {d.strftime('%Y-%m-%d'): float(t or 0) for d, t in rows}
    return jsonify({'year': year, 'month': month, 'by_day': by_day})


@app.route('/api/reports/combined_summary')
@jwt_required()
def reports_combined_summary():
    year = request.args.get('year', type=int, default=datetime.utcnow().year)
    
    # Get income data by month
    income_rows = []
    try:
        income_rows = (
            db.session.query(db.extract('month', Invoice.date).label('month'), db.func.sum(Invoice.total))
            .filter(db.extract('year', Invoice.date) == year)
            .filter(Invoice.type == 'factura')
            .filter(Invoice.paid.is_(True))
            .group_by('month')
            .order_by('month')
            .all()
        )
    except Exception:
        income_rows = db.session.execute(
            text("""
                SELECT CAST(STRFTIME('%m', date) AS INTEGER) AS month, SUM(total)
                FROM invoice
                WHERE type = 'factura' AND paid = 1 AND CAST(STRFTIME('%Y', date) AS INTEGER) = :year
                GROUP BY month
                ORDER BY month
            """), { 'year': year }
        ).fetchall()
    
    # Get expenses data by month
    expenses_rows = []
    try:
        expenses_rows = (
            db.session.query(db.extract('month', Expense.date).label('month'), db.func.sum(Expense.total))
            .filter(db.extract('year', Expense.date) == year)
            .group_by('month')
            .order_by('month')
            .all()
        )
    except Exception:
        expenses_rows = db.session.execute(
            text("""
                SELECT CAST(STRFTIME('%m', date) AS INTEGER) AS month, SUM(total)
                FROM expense
                WHERE CAST(STRFTIME('%Y', date) AS INTEGER) = :year
                GROUP BY month
                ORDER BY month
            """), { 'year': year }
        ).fetchall()
    
    # Create combined data structure
    income_by_month = {int(m): float(t or 0) for m, t in income_rows}
    expenses_by_month = {int(m): float(t or 0) for m, t in expenses_rows}
    
    # Calculate profit (income - expenses) for each month
    profit_by_month = {}
    for month in range(1, 13):
        income = income_by_month.get(month, 0)
        expenses = expenses_by_month.get(month, 0)
        profit_by_month[month] = income - expenses
    
    # Calculate totals
    total_income = sum(income_by_month.values())
    total_expenses = sum(expenses_by_month.values())
    total_profit = total_income - total_expenses
    
    return jsonify({
        'year': year,
        'income_by_month': income_by_month,
        'expenses_by_month': expenses_by_month,
        'profit_by_month': profit_by_month,
        'total_income': total_income,
        'total_expenses': total_expenses,
        'total_profit': total_profit
    })


@app.route('/api/reports/monthly_summary')
@jwt_required()
def reports_monthly_summary():
    year = request.args.get('year', type=int, default=datetime.utcnow().year)
    month = request.args.get('month', type=int, default=datetime.utcnow().month)
    
    # Get income data for the specific month
    income_month = 0
    try:
        result = (
            db.session.query(db.func.sum(Invoice.total))
            .filter(db.extract('year', Invoice.date) == year)
            .filter(db.extract('month', Invoice.date) == month)
            .filter(Invoice.type == 'factura')
            .filter(Invoice.paid.is_(True))
            .scalar()
        )
        income_month = float(result or 0)
    except Exception:
        result = db.session.execute(
            text("""
                SELECT SUM(total)
                FROM invoice
                WHERE type = 'factura'
                  AND paid = 1
                  AND CAST(STRFTIME('%Y', date) AS INTEGER) = :year
                  AND CAST(STRFTIME('%m', date) AS INTEGER) = :month
            """), { 'year': year, 'month': month }
        ).scalar()
        income_month = float(result or 0)
    
    # Get expenses data for the specific month
    expenses_month = 0
    try:
        result = (
            db.session.query(db.func.sum(Expense.total))
            .filter(db.extract('year', Expense.date) == year)
            .filter(db.extract('month', Expense.date) == month)
            .scalar()
        )
        expenses_month = float(result or 0)
    except Exception:
        result = db.session.execute(
            text("""
                SELECT SUM(total)
                FROM expense
                WHERE CAST(STRFTIME('%Y', date) AS INTEGER) = :year
                  AND CAST(STRFTIME('%m', date) AS INTEGER) = :month
            """), { 'year': year, 'month': month }
        ).scalar()
        expenses_month = float(result or 0)
    
    # Calculate profit for the month
    profit_month = income_month - expenses_month
    
    return jsonify({
        'year': year,
        'month': month,
        'income_month': income_month,
        'expenses_month': expenses_month,
        'profit_month': profit_month
    })


def _csv_response(filename: str, content: str) -> Response:
    return Response(
        content,
        mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'}
    )


@app.route('/api/clients/export')
@jwt_required()
@limiter.limit("10 per minute")
def export_clients():
    clients = Client.query.order_by(Client.id).all()
    lines = ['id,name,cif,address,email,phone,iban,created_at']
    for c in clients:
        lines.append(
            f'{c.id},"{c.name}",{c.cif},"{c.address}",{c.email},{c.phone},{c.iban or ""},{(c.created_at or "")}'
        )
    return _csv_response('clientes.csv', '\n'.join(lines))


@app.route('/api/invoices/export')
@jwt_required()
@limiter.limit("10 per minute")
def export_invoices():
    invoices = Invoice.query.order_by(Invoice.id).all()
    lines = ['id,number,date,type,client_id,total,tax_total']
    for inv in invoices:
        lines.append(
            f'{inv.id},{inv.number},{inv.date.isoformat()},{inv.type},{inv.client_id},{inv.total},{inv.tax_total}'
        )
    return _csv_response('facturas.csv', '\n'.join(lines))


@app.route('/api/clients/export_xlsx')
@jwt_required()
@limiter.limit("10 per minute")
def export_clients_xlsx():
    # Lazy import to avoid hard dependency if not used
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = 'Clientes'
    ws.append(['id','name','cif','address','email','phone','iban','created_at'])
    for c in Client.query.order_by(Client.id).all():
        ws.append([c.id, c.name, c.cif, c.address, c.email, c.phone, c.iban or '', c.created_at])
    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return send_file(bio, as_attachment=True, download_name='clientes.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/invoices/export_xlsx')
@jwt_required()
@limiter.limit("10 per minute")
def export_invoices_xlsx():
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = 'Facturas'
    ws.append(['id','number','date','type','client_id','total','tax_total'])
    for inv in Invoice.query.order_by(Invoice.id).all():
        ws.append([inv.id, inv.number, inv.date.isoformat(), inv.type, inv.client_id, inv.total, inv.tax_total])
    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return send_file(bio, as_attachment=True, download_name='facturas.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/expenses', methods=['POST'])
@jwt_required()
@limiter.limit("60 per minute")
def create_expense():
    """Create a new expense record."""
    data = request.get_json(force=True)
    
    # Basic validation
    required_fields = ['date', 'category', 'description', 'supplier', 'base_amount']
    for field in required_fields:
        if not data.get(field):
            return jsonify({'error': f'{field} is required'}), 400
    
    # Validate numeric fields
    try:
        base_amount = float(data['base_amount'])
        if base_amount < 0:
            return jsonify({'error': 'Base amount must be >= 0'}), 400
    except (ValueError, TypeError):
        return jsonify({'error': 'Base amount must be a valid number'}), 400
    
    tax_rate = data.get('tax_rate', 21.0)
    try:
        tax_rate = float(tax_rate)
        if tax_rate < 0 or tax_rate > 100:
            return jsonify({'error': 'Tax rate must be between 0 and 100'}), 400
    except (ValueError, TypeError):
        return jsonify({'error': 'Tax rate must be a valid number'}), 400
    
    # Calculate total if not provided
    if 'total' in data:
        try:
            total = float(data['total'])
            if total < 0:
                return jsonify({'error': 'Total must be >= 0'}), 400
        except (ValueError, TypeError):
            return jsonify({'error': 'Total must be a valid number'}), 400
    else:
        total = round(base_amount * (1 + tax_rate / 100), 2)
    
    # Convert date string to date object
    try:
        date_obj = datetime.strptime(data['date'], '%Y-%m-%d').date()
    except ValueError:
        return jsonify({'error': 'Invalid date format. Use YYYY-MM-DD'}), 400
    
    expense = Expense(
        date=date_obj,
        category=data['category'],
        description=data['description'],
        supplier=data['supplier'],
        base_amount=base_amount,
        tax_rate=tax_rate,
        total=total,
        paid=data.get('paid', False)
    )
    
    db.session.add(expense)
    db.session.commit()
    
    return jsonify({
        'id': expense.id,
        'date': expense.date.isoformat(),
        'category': expense.category,
        'description': expense.description,
        'supplier': expense.supplier,
        'base_amount': expense.base_amount,
        'tax_rate': expense.tax_rate,
        'total': expense.total,
        'paid': expense.paid,
        'created_at': expense.created_at.isoformat()
    }), 201


@app.route('/api/expenses', methods=['GET'])
@jwt_required()
def list_expenses():
    """List expenses with pagination, search and sorting."""
    limit = request.args.get('limit', type=int, default=10)
    offset = request.args.get('offset', type=int, default=0)
    q = request.args.get('q', '')
    sort = request.args.get('sort', 'date')
    dir = request.args.get('dir', 'desc')
    
    # Validate sort field
    allowed_sort = {'date', 'total', 'created_at', 'category', 'description', 'supplier'}
    if sort not in allowed_sort:
        sort = 'date'
    
    # Validate direction
    if dir not in {'asc', 'desc'}:
        dir = 'desc'
    
    query = Expense.query
    
    # Apply search filter
    if q:
        search_term = f'%{q}%'
        query = query.filter(
            db.or_(
                Expense.description.ilike(search_term),
                Expense.supplier.ilike(search_term),
                Expense.category.ilike(search_term)
            )
        )
    
    # Apply sorting (with created_at as secondary sort for consistent ordering)
    sort_field = getattr(Expense, sort)
    if dir == 'desc':
        query = query.order_by(sort_field.desc(), Expense.created_at.desc())
    else:
        query = query.order_by(sort_field.asc(), Expense.created_at.desc())
    
    total = query.count()
    expenses = query.offset(offset).limit(limit).all()
    
    items = []
    for exp in expenses:
        items.append({
            'id': exp.id,
            'date': exp.date.isoformat(),
            'category': exp.category,
            'description': exp.description,
            'supplier': exp.supplier,
            'base_amount': exp.base_amount,
            'tax_rate': exp.tax_rate,
            'total': exp.total,
            'paid': exp.paid,
            'created_at': exp.created_at.isoformat()
        })
    
    return jsonify({'items': items, 'total': total})


@app.route('/api/expenses/<int:expense_id>', methods=['PUT'])
@jwt_required()
def update_expense(expense_id):
    """Update an existing expense record."""
    expense = Expense.query.get_or_404(expense_id)
    data = request.get_json(force=True)
    
    # Update fields with validation
    if 'date' in data:
        try:
            expense.date = datetime.strptime(data['date'], '%Y-%m-%d').date()
        except ValueError:
            return jsonify({'error': 'Invalid date format. Use YYYY-MM-DD'}), 400
    
    if 'category' in data:
        expense.category = data['category']
    if 'description' in data:
        expense.description = data['description']
    if 'supplier' in data:
        expense.supplier = data['supplier']
    if 'paid' in data:
        expense.paid = data['paid']
    
    # Handle numeric fields with validation
    if 'base_amount' in data:
        try:
            base_amount = float(data['base_amount'])
            if base_amount < 0:
                return jsonify({'error': 'Base amount must be >= 0'}), 400
            expense.base_amount = base_amount
        except (ValueError, TypeError):
            return jsonify({'error': 'Base amount must be a valid number'}), 400
    
    if 'tax_rate' in data:
        try:
            tax_rate = float(data['tax_rate'])
            if tax_rate < 0 or tax_rate > 100:
                return jsonify({'error': 'Tax rate must be between 0 and 100'}), 400
            expense.tax_rate = tax_rate
        except (ValueError, TypeError):
            return jsonify({'error': 'Tax rate must be a valid number'}), 400
    
    # Recalculate total if base_amount or tax_rate changed
    if 'base_amount' in data or 'tax_rate' in data:
        expense.total = round(expense.base_amount * (1 + expense.tax_rate / 100), 2)
    elif 'total' in data:
        try:
            total = float(data['total'])
            if total < 0:
                return jsonify({'error': 'Total must be >= 0'}), 400
            expense.total = total
        except (ValueError, TypeError):
            return jsonify({'error': 'Total must be a valid number'}), 400
    
    db.session.commit()
    
    return jsonify({'status': 'ok'})


@app.route('/api/expenses/<int:expense_id>', methods=['DELETE'])
@jwt_required()
def delete_expense(expense_id):
    """Delete an expense record."""
    expense = Expense.query.get_or_404(expense_id)
    db.session.delete(expense)
    db.session.commit()
    return jsonify({'status': 'deleted'})


@app.route('/api/expenses/categories', methods=['GET'])
@jwt_required()
def get_expense_categories():
    """Get unique expense categories for autocomplete."""
    categories = db.session.query(Expense.category).distinct().order_by(Expense.category).all()
    # Filtrar categorías válidas: no vacías, mínimo 2 caracteres, contiene al menos una letra
    valid_categories = []
    for cat in categories:
        if cat[0]:
            trimmed = cat[0].strip()
            # Debe tener al menos 2 caracteres y contener al menos una letra
            if len(trimmed) >= 2 and any(c.isalpha() for c in trimmed):
                valid_categories.append(trimmed)
    return jsonify({'categories': valid_categories})


@app.route('/api/expenses/suppliers', methods=['GET'])
@jwt_required()
def get_expense_suppliers():
    """Get unique expense suppliers for autocomplete."""
    suppliers = db.session.query(Expense.supplier).distinct().order_by(Expense.supplier).all()
    # Filtrar proveedores válidos: no vacíos, mínimo 2 caracteres
    valid_suppliers = []
    for sup in suppliers:
        if sup[0] and len(sup[0].strip()) >= 2:
            valid_suppliers.append(sup[0].strip())
    return jsonify({'suppliers': valid_suppliers})


@app.route('/api/expenses/export_xlsx')
@jwt_required()
@limiter.limit("10 per minute")
def export_expenses_xlsx():
    """Export expenses to XLSX file."""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = 'Gastos'
    ws.append(['id', 'date', 'category', 'description', 'supplier', 'base_amount', 'tax_rate', 'total', 'paid', 'created_at'])
    for exp in Expense.query.order_by(Expense.id).all():
        ws.append([
            exp.id, 
            exp.date.isoformat(), 
            exp.category, 
            exp.description, 
            exp.supplier, 
            exp.base_amount, 
            exp.tax_rate, 
            exp.total, 
            exp.paid, 
            exp.created_at.isoformat()
        ])
    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return send_file(bio, as_attachment=True, download_name='gastos.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/api/invoices/<int:invoice_id>/pdf', methods=['GET'])
@jwt_required()
@limiter.limit("60 per minute")
def invoice_pdf(invoice_id):
    """Generate a PDF for a given invoice."""
    invoice = Invoice.query.get_or_404(invoice_id)
    client = invoice.client
    company = CompanyConfig.query.first() or _company_from_env()
    items = invoice.items
    # Build absolute file URI for logo to avoid network issues in wkhtmltopdf
    # Use a proper file URI (file:///C:/...) and forward slashes
    logo_path = os.path.join(STATIC_FOLDER, 'logo_invoice.png')
    try:
        logo_uri = Path(logo_path).resolve().as_uri()
    except Exception:
        logo_uri = 'file:///' + logo_path.replace('\\', '/')
    rendered = render_template(
        'invoice_template.html',
        invoice=invoice,
        client=client,
        company=company,
        items=items,
        logo_uri=logo_uri,
        company_address_line=_compose_company_address(company),
    )
    filename = f"{invoice.type}_{invoice.number}.pdf"
    file_path = os.path.join(DOWNLOAD_FOLDER, filename)
    # Usar solo wkhtmltopdf para consistencia dev/prod
    pdf_bytes = None
    if pdfkit is not None:
        try:
            options = {
                'enable-local-file-access': None,
                'page-size': 'A4',
                'margin-top': '0.75in',
                'margin-right': '0.75in',
                'margin-bottom': '0.75in',
                'margin-left': '0.75in',
                'encoding': 'UTF-8',
                'no-outline': None,
                'print-media-type': None,
                'disable-smart-shrinking': None,
                'zoom': 1.0,
                'dpi': 96,
                'image-quality': 94,
                'minimum-font-size': 12,
                'enable-plugins': None,
                'load-error-handling': 'ignore',
                'load-media-error-handling': 'ignore'
            }
            cfg = _resolve_pdfkit_configuration()
            pdf_bytes = pdfkit.from_string(rendered, False, options=options, configuration=cfg)
        except Exception:
            pdf_bytes = None
    
    # Fallback solo si wkhtmltopdf falla
    if pdf_bytes is None:
        pdf_bytes = _generate_pdf_fallback(invoice, client, company, items)
    # Save the PDF to the downloads folder
    with open(file_path, 'wb') as f:
        f.write(pdf_bytes)
    pdf_io = io.BytesIO(pdf_bytes)
    return send_file(pdf_io, mimetype='application/pdf', as_attachment=True,
                     download_name=filename)


@app.route('/health', methods=['GET'])
def health():
    """Sonda de salud simple para monitoreo/compose."""
    return jsonify({
        'status': 'ok',
    'version': os.getenv('APP_VERSION', 'dev'),
        'pdf_engine': 'wkhtmltopdf' if pdfkit else 'reportlab_fallback',
        'database': app.config.get('SQLALCHEMY_DATABASE_URI', '')
    })


@app.route('/static/uploads/products/<path:filename>', methods=['GET'])
def serve_product_image(filename):
    """Sirve imágenes de productos sin autenticación (archivos públicos)."""
    from flask import send_from_directory
    products_folder = os.path.join(STATIC_FOLDER, 'uploads', 'products')
    return send_from_directory(products_folder, filename)


# -------------------------------------
# Auth
# -------------------------------------

@app.post('/api/auth/register')
@jwt_required(optional=True)
@limiter.limit("5 per minute")
def register():
    """Crea el primer usuario sin token; posteriores requieren token."""
    identity = get_jwt_identity()
    data = request.get_json(force=True)
    username = data.get('username')
    password = data.get('password')
    if not username or not password:
        return jsonify({'error': 'username y password requeridos'}), 400
    existing_count = User.query.count()
    if existing_count > 0 and not identity:
        return jsonify({'error': 'No autorizado'}), 401
    if User.query.filter_by(username=username).first():
        return jsonify({'error': 'Usuario ya existe'}), 409
    user = User(username=username, password_hash=generate_password_hash(password))
    db.session.add(user)
    db.session.commit()
    return jsonify({'status': 'ok'}), 201


@app.post('/api/auth/login')
@limiter.limit("10 per minute")
def login():
    """
    Authenticate user and return JWT access token.
    
    Validates username/password combination and returns JWT token for API access.
    Token expires after 30 days.
    
    Returns:
        JSON: Access token on success or error message on failure
    """
    try:
        payload = LoginRequest.model_validate(request.get_json(force=True))
    except PydValidationError as e:
        return jsonify({"error": e.errors()[0]['msg'] if e.errors() else 'invalid payload', "code": 400}), 400
    username = payload.username
    password = payload.password
    user = User.query.filter_by(username=username).first()
    if not user or not check_password_hash(user.password_hash, password):
        return jsonify({'error': 'Credenciales inválidas', 'code': 401}), 401
    token = create_access_token(identity=username)
    resp = jsonify({'access_token': token})
    # Solo establecer cookies si no están deshabilitadas explícitamente
    if not disable_jwt_cookies:
        try:
            set_access_cookies(resp, token)
        except Exception:
            pass
    return resp


# -------------------------------------
# OpenAPI docs
# -------------------------------------

@app.get('/openapi.json')
def openapi_json():
    base = request.host_url.rstrip('/')
    return jsonify(get_openapi_spec(base)), 200


@app.get('/apidocs')
def apidocs():
    # Minimal Redoc page served inline (uses CDN)
    html = f"""
<!doctype html>
<html>
  <head>
    <meta charset=\"utf-8\"/>
    <title>Nioxtec Facturer API Docs</title>
    <style> body {{ margin:0; padding:0; }} </style>
  </head>
  <body>
    <redoc spec-url=\"/openapi.json\"></redoc>
    <script src=\"https://cdn.jsdelivr.net/npm/redoc/bundles/redoc.standalone.js\"></script>
  </body>
</html>
"""
    return Response(html, mimetype='text/html')


# -------------------------------------
# Contract Generation
# -------------------------------------

@app.post('/api/auth/logout')
def logout():
    """Borra cookies JWT en el navegador (logout)."""
    resp = jsonify({'status': 'ok'})
    unset_jwt_cookies(resp)
    return resp

@app.post('/api/contracts/generate-pdf')
@jwt_required()
@limiter.limit("25 per minute")
def generate_contract_pdf():
    """
    Generate PDF from contract content.
    
    Accepts DOCX template ID and form data, fills the template and generates a PDF.
    Optionally saves the PDF as client document if client_id is provided.
    
    Returns:
        JSON: Generated filename on success
    """
    data = request.get_json(force=True)
    template_id = data.get('template_id')
    form_data = data.get('form_data', {})
    filename = data.get('filename', 'contract.pdf')
    client_id = data.get('client_id')  # Optional: if provided, save as client document
    
    if not template_id:
        return jsonify({'error': 'Template ID is required'})
    
    try:
        # Get template filename
        template_filename = None
        if template_id == 'compraventa':
            template_filename = 'Contrato_Compraventa_Plazos_NIOXTEC_v5.docx'
        elif template_id == 'renting':
            template_filename = 'Plantilla_Contrato_Renting_Firma_Datos_v2.docx'
        else:
            return jsonify({'error': 'Invalid template ID'}), 400
        
        # Load and fill DOCX template
        template_path = os.path.join(STATIC_FOLDER, 'contracts', 'templates', template_filename)
        if not os.path.exists(template_path):
            return jsonify({'error': 'Template file not found'}), 404
        
        # Load DOCX document
        doc = Document(template_path)
        
        # Get original placeholders from template
        template_info = extract_placeholders_from_docx(template_filename)
        if 'error' in template_info:
            return jsonify({'error': template_info['error']}), 400
        
        original_tokens = template_info.get('original_tokens', {})
        
        # Mapping from document placeholders to form data keys
        placeholder_mapping = {
            # Compraventa template - usar solo las claves principales
            'Nombre completo del cliente': 'nombre_completo_del_cliente',
            'DNI DEL CLIENTE': 'numero',
            'Dirección del cliente': 'direccion',
            'Teléfono del cliente': 'telefono',
            'Correo del cliente': 'correo',
            'Modelo del producto': 'modelo',
            'Pulgadas del producto': 'pulgadas',
            'Número de serie del producto': 'numero_serie',
            'importe total en euros, IVA incluido': 'importe_total_en_euros_iva_incluido',
            'número de plazos': 'numero_de_plazos',
            'importe de cada cuota': 'importe_de_cada_cuota',
            'Tabla de interes': 'tabla_de_interes',
            
            # Campos duplicados - mapear a las mismas claves principales
            'Nombre del comprador': 'nombre_completo_del_cliente',  # = Nombre completo del cliente
            'Dni del comprador': 'numero',  # = DNI DEL CLIENTE
            'Modelo': 'modelo',  # = Modelo del producto
            'Pulgadas': 'pulgadas',  # = Pulgadas del producto
            'Número de Serie': 'numero_serie',  # = Número de serie del producto
            
            # Renting template
            'Nombre de la empresa o persona': 'nombre_de_la_empresa_o_persona',
            'Número': 'numero',
            'Dirección': 'direccion',
            'Nombre representante': 'nombre_representante',
            'Cargo': 'cargo',
            'Teléfono': 'telefono',
            'Correo': 'correo',
            'Marca': 'marca',
            'importe en euros': 'importe_en_euros',
            'plataforma de pago': 'plataforma_de_pago',
            'IBAN': 'iban',
            'importe ajustado': 'importe_ajustado',
            # Fecha actual en formato DD-MM-AAAA (nueva etiqueta en plantillas)
            'FECHA FORMATO DD-MM-AAAA': 'fecha_formato_dd_mm_aaaa',
        }

        # Asegurar fecha actual si la plantilla la requiere y no viene en el formulario
        form_data = dict(form_data or {})
        if 'fecha_formato_dd_mm_aaaa' not in form_data:
            try:
                form_data['fecha_formato_dd_mm_aaaa'] = datetime.now().strftime('%d-%m-%Y')
            except Exception:
                form_data['fecha_formato_dd_mm_aaaa'] = ''
        
        # Interest table mapping based on number of installments
        def get_interest_text(num_plazos):
            if 0 <= num_plazos <= 3:
                return "Sin intereses"
            elif 4 <= num_plazos <= 6:
                return "5% interés mensual"
            elif 7 <= num_plazos <= 12:
                return "10% interés mensual"
            elif 13 <= num_plazos <= 18:
                return "20% interés mensual"
            elif 19 <= num_plazos <= 24:
                return "30% interés mensual"
            else:
                return "Sin intereses"  # Por defecto
        
        # Calculate interest based on number of installments
        if template_id == 'compraventa' and 'numero_de_plazos' in form_data:
            num_plazos = int(form_data.get('numero_de_plazos', 3))
            interest_text = get_interest_text(num_plazos)
            form_data['tabla_de_interes'] = interest_text
        
        # Debug logging
        app.logger.info(f"Template placeholders: {list(original_tokens.keys())}")
        app.logger.info(f"Form data keys: {list(form_data.keys())}")
        app.logger.info(f"Form data: {form_data}")
        app.logger.info(f"Placeholder mapping: {placeholder_mapping}")
        
        # Log which placeholders will be replaced
        replacements_made = []
        
        # Fill placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder_name, placeholder_token in original_tokens.items():
                # Map placeholder to form data key
                form_key = placeholder_mapping.get(placeholder_name, placeholder_name)
                
                if form_key in form_data:
                    value = str(form_data[form_key])
                    if placeholder_token in paragraph.text:
                        app.logger.info(f"Replacing {placeholder_token} with {value} in paragraph (key: {form_key})")
                        paragraph.text = paragraph.text.replace(placeholder_token, value)
                        replacements_made.append(f"{placeholder_name} -> {value}")
        
        # Fill placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder_name, placeholder_token in original_tokens.items():
                            # Map placeholder to form data key
                            form_key = placeholder_mapping.get(placeholder_name, placeholder_name)
                            
                            if form_key in form_data:
                                value = str(form_data[form_key])
                                if placeholder_token in paragraph.text:
                                    app.logger.info(f"Replacing {placeholder_token} with {value} in table cell (key: {form_key})")
                                    paragraph.text = paragraph.text.replace(placeholder_token, value)
                                    replacements_made.append(f"{placeholder_name} -> {value}")
        
        app.logger.info(f"Total replacements made: {len(replacements_made)}")
        app.logger.info(f"Replacements: {replacements_made}")
        
        # Save filled DOCX temporarily
        temp_docx_path = os.path.join(DOWNLOAD_FOLDER, f'temp_{secure_filename(filename)}.docx')
        doc.save(temp_docx_path)
        
        # Convert DOCX to PDF using wkhtmltopdf
        if pdfkit is not None:
            # Convert DOCX to HTML first (simple approach)
            html_content = _docx_to_html(temp_docx_path)
            
            # Create contract template HTML
            contract_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{
                        font-family: "Cambria", "Times New Roman", serif;
                        line-height: 1.4;
                        margin: 1.5cm 1.2cm 1.5cm 1.2cm;
                        font-size: 11pt;
                        color: #333;
                    }}
                    .header {{
                        display: flex;
                        justify-content: center;
                        align-items: center;
                        text-align: center;
                        margin-bottom: 1.5em;
                        page-break-inside: avoid;
                        width: 100%;
                    }}
                    .header img {{
                        max-height: 70px;
                        max-width: 160px;
                        object-fit: contain;
                        margin: 0 auto;
                        display: block;
                    }}
                    /* Párrafos con mejor espaciado */
                    p {{
                        margin: 0.5em 0;
                        text-align: justify;
                    }}
                    /* Negritas destacadas */
                    strong {{
                        font-weight: bold;
                        color: #222;
                    }}
                    /* Jerarquía visual mejorada para títulos */
                    h1, h2, h3 {{
                        color: #65AAC3;
                        font-family: "Cambria", "Georgia", "Times New Roman", serif;
                        font-weight: bold;
                        page-break-after: avoid;
                        line-height: 1.2;
                    }}
                    h1 {{ 
                        font-size: 14pt;
                        margin-top: 0.5em;
                        margin-bottom: 1.2em;
                        letter-spacing: 0.5px;
                        font-family: "Cambria", "Georgia", "Times New Roman", serif;
                    }}
                    h2 {{ 
                        font-size: 12pt;
                        margin-top: 2em;
                        margin-bottom: 0.8em;
                        padding-top: 0.5em;
                        border-top: 2px solid #65AAC3;
                        font-family: "Cambria", "Georgia", "Times New Roman", serif;
                    }}
                    h3 {{ 
                        font-size: 10pt;
                        margin-top: 1.2em;
                        margin-bottom: 0.6em;
                        font-family: "Cambria", "Georgia", "Times New Roman", serif;
                    }}
                    /* Estilos específicos para títulos del contrato */
                    .contract-title {{
                        font-size: 16pt;
                        color: #65AAC3;
                        font-weight: bold;
                        text-align: center;
                        margin-bottom: 1.5em;
                        text-transform: uppercase;
                        letter-spacing: 0.5px;
                        font-family: "Cambria", "Georgia", "Times New Roman", serif;
                    }}
                    .section-title {{
                        font-size: 14pt;
                        color: #65AAC3;
                        font-weight: bold;
                        margin-top: 2.5em;
                        margin-bottom: 1em;
                        padding-top: 0.8em;
                        padding-bottom: 0.3em;
                        border-top: 2px solid #65AAC3;
                        border-bottom: 1px solid #65AAC3;
                        font-family: "Cambria", "Georgia", "Times New Roman", serif;
                    }}
                    .subsection-title {{
                        font-size: 12pt;
                        color: #65AAC3;
                        font-weight: bold;
                        margin-top: 1.5em;
                        margin-bottom: 0.7em;
                        font-family: "Cambria", "Georgia", "Times New Roman", serif;
                    }}
                    /* Tablas mejoradas */
                    table {{
                        width: 100%;
                        border-collapse: collapse;
                        margin: 1.2em 0;
                        font-family: "Cambria", "Times New Roman", serif;
                        font-size: 10.5pt;
                        page-break-inside: avoid;
                    }}
                    th, td {{
                        border: 1px solid #ccc;
                        padding: 10px 12px;
                        text-align: left;
                        font-family: "Cambria", "Times New Roman", serif;
                    }}
                    th {{
                        background-color: #65AAC3;
                        color: white;
                        font-weight: bold;
                        font-size: 11pt;
                    }}
                    td {{
                        background-color: #fafafa;
                    }}
                    tr:nth-child(even) td {{
                        background-color: #f5f5f5;
                    }}
                    /* Sección de firma */
                    .signature-section {{
                        margin-top: 3em;
                        page-break-before: auto;
                        page-break-inside: avoid;
                    }}
                    .signature-line {{
                        border-top: 2px solid #333;
                        margin-top: 2.5em;
                        padding-top: 0.8em;
                        text-align: center;
                        font-weight: bold;
                    }}
                </style>
            </head>
            <body>
                <div class="header">
                    <img src="{Path(os.path.join(STATIC_FOLDER, 'contracts', 'images', 'header-right.png')).resolve().as_uri()}" alt="NIOXTEC Logo" />
                </div>
                {html_content}
            </body>
            </html>
            """
            
            options = {
                'enable-local-file-access': None,
                'page-size': 'A4',
                'margin-top': '0.75cm',
                'margin-right': '1.2cm',
                'margin-bottom': '0.75cm',
                'margin-left': '1.2cm',
                'encoding': 'UTF-8',
                'no-outline': None,
                'print-media-type': None,
                'disable-smart-shrinking': None,
                'zoom': 1.0,
                'dpi': 96,
                'minimum-font-size': 10
            }
            cfg = _resolve_pdfkit_configuration()
            pdf_bytes = pdfkit.from_string(contract_html, False, options=options, configuration=cfg)
        else:
            # Fallback to reportlab
            pdf_bytes = _generate_contract_pdf_fallback(_docx_to_text(temp_docx_path))
        
        # Clean up temporary DOCX file
        if os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)
        
        # Save PDF to downloads folder
        safe_filename = secure_filename(filename)
        file_path = os.path.join(DOWNLOAD_FOLDER, safe_filename)
        with open(file_path, 'wb') as f:
            f.write(pdf_bytes)
        
        # If client_id is provided, also save as client document
        document_saved = False
        if client_id:
            try:
                # Verify client exists
                client = Client.query.get(client_id)
                if client:
                    # Check if document with same filename already exists for this client
                    existing_doc = ClientDocument.query.filter_by(
                        client_id=client_id,
                        filename=safe_filename
                    ).first()
                    
                    if not existing_doc:
                        # Save to client documents folder
                        base_dir = _client_upload_dir(client_id)
                        documents_dir = os.path.join(base_dir, 'documents')
                        os.makedirs(documents_dir, exist_ok=True)
                        
                        unique_name = f"{uuid4().hex}_{safe_filename}"
                        stored_rel = os.path.join(str(client_id), 'documents', unique_name)
                        stored_abs = os.path.join(UPLOADS_ROOT, stored_rel)
                        
                        with open(stored_abs, 'wb') as f:
                            f.write(pdf_bytes)
                        
                        # Save document record to database
                        doc_record = ClientDocument(
                            client_id=client_id,
                            category='document',
                            filename=safe_filename,
                            stored_path=stored_rel,
                            content_type='application/pdf',
                            size_bytes=len(pdf_bytes),
                        )
                        db.session.add(doc_record)
                        db.session.commit()
                        document_saved = True
                        app.logger.info(f"Contract PDF auto-saved as client document: {safe_filename} for client {client_id}")
            except Exception as save_error:
                app.logger.error(f"Error auto-saving contract as client document: {save_error}")
                # Don't fail the whole request, just log the error
        
        return jsonify({
            'filename': safe_filename,
            'document_saved': document_saved
        })
        
    except Exception as e:
        app.logger.error(f"Error generating contract PDF: {e}")
        return jsonify({'error': 'Error generating PDF'}), 500


@app.get('/api/contracts/download/<filename>')
@jwt_required()
def download_contract_pdf(filename):
    """
    Download generated contract PDF.
    
    Args:
        filename: Name of the PDF file to download
        
    Returns:
        PDF file as attachment
    """
    try:
        file_path = os.path.join(DOWNLOAD_FOLDER, secure_filename(filename))
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404
        
        return send_file(file_path, mimetype='application/pdf', as_attachment=True)
        
    except Exception as e:
        app.logger.error(f"Error downloading contract PDF: {e}")
        return jsonify({'error': 'Error downloading file'}), 500


@app.post('/api/contracts/save-as-document')
@jwt_required()
@limiter.limit("10 per minute")
def save_contract_as_document():
    """
    Save generated contract PDF as client document.
    
    Accepts template ID, form data, filename and client ID, generates PDF and saves it as client document.
    
    Returns:
        JSON: Saved document info on success
    """
    data = request.get_json(force=True)
    template_id = data.get('template_id')
    form_data = data.get('form_data', {})
    filename = data.get('filename', 'contract.pdf')
    client_id = data.get('client_id')
    
    if not template_id:
        return jsonify({'error': 'Template ID is required'}), 400
    
    if not client_id:
        return jsonify({'error': 'Client ID is required'}), 400
    
    # Verify client exists
    client = Client.query.get(client_id)
    if not client:
        return jsonify({'error': 'Client not found'}), 404
    
    try:
        # Get template filename
        template_filename = None
        if template_id == 'compraventa':
            template_filename = 'Contrato_Compraventa_Plazos_NIOXTEC_v5.docx'
        elif template_id == 'renting':
            template_filename = 'Plantilla_Contrato_Renting_Firma_Datos_v2.docx'
        else:
            return jsonify({'error': 'Invalid template ID'}), 400
        
        # Load and fill DOCX template
        template_path = os.path.join(STATIC_FOLDER, 'contracts', 'templates', template_filename)
        if not os.path.exists(template_path):
            return jsonify({'error': 'Template file not found'}), 404
        
        # Load DOCX document
        doc = Document(template_path)
        
        # Get original placeholders from template
        template_info = extract_placeholders_from_docx(template_filename)
        if 'error' in template_info:
            return jsonify({'error': template_info['error']}), 400
        
        original_tokens = template_info.get('original_tokens', {})
        
        # Mapping from document placeholders to form data keys
        placeholder_mapping = {
            # Compraventa template - usar solo las claves principales
            'Nombre completo del cliente': 'nombre_completo_del_cliente',
            'DNI DEL CLIENTE': 'numero',
            'Dirección del cliente': 'direccion',
            'Teléfono del cliente': 'telefono',
            'Correo del cliente': 'correo',
            'Modelo del producto': 'modelo',
            'Pulgadas del producto': 'pulgadas',
            'Número de serie del producto': 'numero_serie',
            'importe total en euros, IVA incluido': 'importe_total_en_euros_iva_incluido',
            'número de plazos': 'numero_de_plazos',
            'importe de cada cuota': 'importe_de_cada_cuota',
            'Tabla de interes': 'tabla_de_interes',
            
            # Campos duplicados - mapear a las mismas claves principales
            'Nombre del comprador': 'nombre_completo_del_cliente',  # = Nombre completo del cliente
            'Dni del comprador': 'numero',  # = DNI DEL CLIENTE
            'Modelo': 'modelo',  # = Modelo del producto
            'Pulgadas': 'pulgadas',  # = Pulgadas del producto
            'Número de Serie': 'numero_serie',  # = Número de serie del producto
            
            # Renting template
            'Nombre de la empresa o persona': 'nombre_de_la_empresa_o_persona',
            'Número': 'numero',
            'Dirección': 'direccion',
            'Nombre representante': 'nombre_representante',
            'Cargo': 'cargo',
            'Teléfono': 'telefono',
            'Correo': 'correo',
            'Marca': 'marca',
            'importe en euros': 'importe_en_euros',
            'plataforma de pago': 'plataforma_de_pago',
            'IBAN': 'iban',
            'importe ajustado': 'importe_ajustado',
            # Fecha actual en formato DD-MM-AAAA (nueva etiqueta en plantillas)
            'FECHA FORMATO DD-MM-AAAA': 'fecha_formato_dd_mm_aaaa',
        }

        # Asegurar fecha actual si la plantilla la requiere y no viene en el formulario
        form_data = dict(form_data or {})
        if 'fecha_formato_dd_mm_aaaa' not in form_data:
            try:
                form_data['fecha_formato_dd_mm_aaaa'] = datetime.now().strftime('%d-%m-%Y')
            except Exception:
                form_data['fecha_formato_dd_mm_aaaa'] = ''
        
        # Interest table mapping based on number of installments
        def get_interest_text(num_plazos):
            if 0 <= num_plazos <= 3:
                return "Sin intereses"
            elif 4 <= num_plazos <= 6:
                return "5% interés mensual"
            elif 7 <= num_plazos <= 12:
                return "10% interés mensual"
            elif 13 <= num_plazos <= 18:
                return "20% interés mensual"
            elif 19 <= num_plazos <= 24:
                return "25% interés mensual"
            else:
                return "30% interés mensual"
        
        # Fill placeholders in document
        for paragraph in doc.paragraphs:
            for placeholder, form_key in placeholder_mapping.items():
                if placeholder in original_tokens:
                    original_token = original_tokens[placeholder]
                    value = form_data.get(form_key, '')
                    
                    # Special handling for interest table
                    if placeholder == 'Tabla de interes':
                        num_plazos = form_data.get('numero_de_plazos', 0)
                        try:
                            num_plazos = int(num_plazos)
                            value = get_interest_text(num_plazos)
                        except (ValueError, TypeError):
                            value = "Sin intereses"
                    
                    # Replace placeholder with value
                    if original_token in paragraph.text:
                        paragraph.text = paragraph.text.replace(original_token, str(value))
        
        # Fill placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, form_key in placeholder_mapping.items():
                            if placeholder in original_tokens:
                                original_token = original_tokens[placeholder]
                                value = form_data.get(form_key, '')
                                
                                # Special handling for interest table
                                if placeholder == 'Tabla de interes':
                                    num_plazos = form_data.get('numero_de_plazos', 0)
                                    try:
                                        num_plazos = int(num_plazos)
                                        value = get_interest_text(num_plazos)
                                    except (ValueError, TypeError):
                                        value = "Sin intereses"
                                
                                # Replace placeholder with value
                                if original_token in paragraph.text:
                                    paragraph.text = paragraph.text.replace(original_token, str(value))
        
        # Save filled DOCX temporarily
        temp_docx_path = os.path.join(DOWNLOAD_FOLDER, f"temp_{uuid4().hex}.docx")
        doc.save(temp_docx_path)
        
        # Generate PDF from filled DOCX
        cfg = _resolve_pdfkit_configuration()
        if pdfkit and cfg:
            # Convert DOCX to HTML first
            contract_html = _docx_to_html(temp_docx_path)
            
            # PDF options
            options = {
                'page-size': 'A4',
                'margin-top': '0.75in',
                'margin-right': '0.75in',
                'margin-bottom': '0.75in',
                'margin-left': '0.75in',
                'encoding': "UTF-8",
                'no-outline': None,
                'enable-local-file-access': None
            }
            
            pdf_bytes = pdfkit.from_string(contract_html, False, options=options, configuration=cfg)
        else:
            # Fallback to reportlab
            pdf_bytes = _generate_contract_pdf_fallback(_docx_to_text(temp_docx_path))
        
        # Clean up temporary DOCX
        try:
            os.remove(temp_docx_path)
        except:
            pass
        
        # Save PDF to client documents folder
        safe_filename = secure_filename(filename)
        
        # Check if document with same filename already exists for this client
        existing_doc = ClientDocument.query.filter_by(
            client_id=client_id, 
            filename=safe_filename
        ).first()
        
        if existing_doc:
            return jsonify({'error': f'Ya existe un documento con el nombre "{safe_filename}" para este cliente'}), 409
        
        base_dir = _client_upload_dir(client_id)
        documents_dir = os.path.join(base_dir, 'documents')
        os.makedirs(documents_dir, exist_ok=True)
        
        unique_name = f"{uuid4().hex}_{safe_filename}"
        stored_rel = os.path.join(str(client_id), 'documents', unique_name)
        stored_abs = os.path.join(UPLOADS_ROOT, stored_rel)
        
        with open(stored_abs, 'wb') as f:
            f.write(pdf_bytes)
        
        # Save document record to database
        doc_record = ClientDocument(
            client_id=client_id,
            category='document',
            filename=safe_filename,
            stored_path=stored_rel,
            content_type='application/pdf',
            size_bytes=len(pdf_bytes),
        )
        db.session.add(doc_record)
        db.session.commit()
        
        return jsonify({
            'id': doc_record.id,
            'filename': safe_filename,
            'size_bytes': len(pdf_bytes),
            'uploaded_at': doc_record.uploaded_at.isoformat()
        })
        
    except Exception as e:
        app.logger.error(f"Error saving contract as document: {e}")
        return jsonify({'error': 'Error saving contract as document'}), 500


@app.get('/static/contracts/images/<filename>')
def serve_contract_image(filename):
    """
    Serve contract images for PDF generation.
    
    Args:
        filename: Name of the image file
        
    Returns:
        Image file
    """
    try:
        image_path = os.path.join('static', 'contracts', 'images', secure_filename(filename))
        if not os.path.exists(image_path):
            return jsonify({'error': 'Image not found'}), 404
        
        return send_file(image_path)
        
    except Exception as e:
        app.logger.error(f"Error serving contract image: {e}")
        return jsonify({'error': 'Error serving image'}), 500


def slugify(text):
    """Convert text to slug format (lowercase, no accents, spaces to underscores)."""
    # Normalize unicode characters
    text = unicodedata.normalize('NFD', text)
    # Remove accents
    text = ''.join(c for c in text if not unicodedata.combining(c))
    # Convert to lowercase and replace spaces with underscores
    text = re.sub(r'[^\w\s-]', '', text.lower())
    text = re.sub(r'[-\s]+', '_', text)
    return text.strip('_')

def extract_placeholders_from_docx(filename):
    """Extract placeholders from DOCX file."""
    try:
        # Path to the template file
        template_path = os.path.join(STATIC_FOLDER, 'contracts', 'templates', filename)
        
        if not os.path.exists(template_path):
            return {'error': f'Template file not found: {filename}'}
        
        doc = Document(template_path)
        placeholders = set()
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Find all placeholders in the paragraph
            matches = re.findall(r'\[(.+?)\]', text)
            for match in matches:
                placeholders.add(match.strip())
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    # Find all placeholders in the cell
                    matches = re.findall(r'\[(.+?)\]', text)
                    for match in matches:
                        placeholders.add(match.strip())
        
        return {
            'placeholders': list(placeholders),
            'original_tokens': {placeholder: f'[{placeholder}]' for placeholder in placeholders}
        }
        
    except Exception as e:
        return {'error': f'Error processing DOCX: {str(e)}'}

def _format_paragraph_with_bold(paragraph):
    """
    Format paragraph text preserving bold formatting from DOCX.
    Returns HTML with <strong> tags for bold text.
    Handles bold that can be: True (explicit), False (explicit), or None (inherited from style).
    """
    html_parts = []
    
    # Check if paragraph style has bold by default
    paragraph_has_bold_style = False
    try:
        if paragraph.style.font.bold:
            paragraph_has_bold_style = True
    except:
        pass
    
    for run in paragraph.runs:
        text = run.text
        if text:
            # Replace newlines with <br>
            text = text.replace('\n', '<br>')
            
            # Determine if this run should be bold
            is_bold = False
            
            if run.bold is True:
                # Explicitly bold
                is_bold = True
            elif run.bold is None and paragraph_has_bold_style:
                # Inherits bold from paragraph style
                is_bold = True
            # If run.bold is False, is_bold stays False (explicitly not bold)
            
            if is_bold:
                html_parts.append(f'<strong>{text}</strong>')
            else:
                html_parts.append(text)
    
    return ''.join(html_parts) if html_parts else paragraph.text.replace('\n', '<br>')

def _docx_to_html(docx_path):
    """Convert DOCX to HTML for PDF generation."""
    doc = Document(docx_path)
    html_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text = paragraph.text.strip()
            
            # Check if paragraph has a heading style
            is_heading_3 = paragraph.style.name in ['Heading 3', 'Título 3', 'Heading3', 'Titulo 3']
            
            # Debug logging
            app.logger.info(f"Processing paragraph: '{text}' (length: {len(text)}, style: {paragraph.style.name})")
            
            # Detectar títulos basándose en el contenido y formato
            if text.upper() == "CONTRATO DE COMPRAVENTA A PLAZOS SIN INTERESES":
                # Título principal del contrato de compraventa
                app.logger.info(f"Detected main title: {text}")
                html_parts.append(f'<h1 class="contract-title" style="font-size: 14pt; color: #65AAC3; font-weight: bold; text-align: center; margin-bottom: 0.5em;">{text}</h1>')
            elif text.upper() == "CONTRATO DE RENTING DE PANTALLA PUBLICITARIA":
                # Título principal del contrato de renting
                app.logger.info(f"Detected main title: {text}")
                html_parts.append(f'<h1 class="contract-title" style="font-size: 14pt; color: #65AAC3; font-weight: bold; text-align: center; margin-bottom: 1.5em; text-transform: uppercase; letter-spacing: 0.5px;">{text}</h1>')
            elif text.upper() == "PARTES INTERVINIENTES":
                # Sección principal del contrato de compraventa
                app.logger.info(f"Detected section title: {text}")
                html_parts.append(f'<h2 class="section-title" style="font-size: 13pt; color: #65AAC3; font-weight: bold; margin-top: 2.5em; margin-bottom: 1em; padding-top: 0.8em; padding-bottom: 0.3em; border-top: 2px solid #65AAC3; border-bottom: 1px solid #65AAC3;">{text}</h2>')
            elif text.upper() in ["CLAUSULAS", "CLÁUSULAS"]:
                # Sección principal del contrato de renting
                app.logger.info(f"Detected section title: {text}")
                html_parts.append(f'<h2 class="section-title" style="font-size: 13pt; color: #65AAC3; font-weight: bold; margin-top: 2.5em; margin-bottom: 1em; padding-top: 0.8em; padding-bottom: 0.3em; border-top: 2px solid #65AAC3; border-bottom: 1px solid #65AAC3;">{text}</h2>')
            elif text.upper() == "ACEPTACIÓN DEL CONTRATO":
                # Sección principal del contrato de renting
                app.logger.info(f"Detected section title: {text}")
                html_parts.append(f'<h2 class="section-title" style="font-size: 13pt; color: #65AAC3; font-weight: bold; margin-top: 2.5em; margin-bottom: 1em; padding-top: 0.8em; padding-bottom: 0.3em; border-top: 2px solid #65AAC3; border-bottom: 1px solid #65AAC3;">{text}</h2>')
            elif text.upper() in ["VENDEDOR", "COMPRADOR", "OBJETO DEL CONTRATO", "GARANTÍA", "IMPAGO", "PROTECCIÓN DE DATOS", "JURISDICCIÓN", "ENTREGA"]:
                # Subsecciones del contrato de compraventa
                # Verificar que sea un título independiente (no parte de una frase)
                if len(text.split()) <= 3 and not any(char in text for char in ['.', ',', ':', ';']):
                    app.logger.info(f"Detected subsection title: {text}")
                    html_parts.append(f'<h3 class="subsection-title" style="font-size: 12pt; color: #65AAC3; font-weight: bold; margin-top: 1.5em; margin-bottom: 0.7em;">{text}</h3>')
                else:
                    # Es parte de una frase, tratarlo como texto normal
                    app.logger.info(f"Title word found in sentence, treating as normal text: {text}")
                    text = text.replace('\n', '<br>')
                    html_parts.append(f'<p style="margin: 0.5em 0; text-align: justify;">{text}</p>')
            elif text.upper() in ["1. OBJETO DEL CONTRATO", "2. DURACIÓN MÍNIMA DEL RENTING", "3. CUOTA DE RENTING Y FORMA DE PAGO", "4. CESIÓN DE PROPIEDAD", "5. USO, INSTALACIÓN Y CONTENIDOS", "6. SERVICIO TÉCNICO Y SOPORTE", "7. RESPONSABILIDAD Y BUENAS PRÁCTICAS", "8. FORMA DE PAGO Y AUTORIZACIÓN SEPA", "9. CANCELACIÓN ANTICIPADA", "10. JURISDICCIÓN"]:
                # Subsecciones del contrato de renting (títulos antiguos)
                app.logger.info(f"Detected subsection title: {text}")
                html_parts.append(f'<h3 class="subsection-title" style="font-size: 12pt; color: #65AAC3; font-weight: bold; margin-top: 1.5em; margin-bottom: 0.7em;">{text}</h3>')
            # Detectar títulos H3 por estilo de Word o por patrón de texto
            elif is_heading_3 or (text.startswith("3.") or text.startswith("4.") or 
                  text.startswith("5. Servicio técnico") or text.startswith("6. Responsabilidad") or 
                  text.startswith("7. Cancelación") or text.startswith("8. Jurisdicción") or
                  text == "Calendario de recobro"):
                # Nuevas subsecciones del contrato de renting H3
                app.logger.info(f"Detected new H3 subsection title: {text} (is_heading_3={is_heading_3})")
                html_parts.append(f'<h3 class="subsection-title" style="font-size: 12pt; color: #65AAC3; font-weight: bold; margin-top: 1.5em; margin-bottom: 0.7em;">{text}</h3>')
            else:
                # Texto normal - detectar negritas en runs del párrafo
                formatted_text = _format_paragraph_with_bold(paragraph)
                html_parts.append(f'<p style="margin: 0.5em 0; text-align: justify;">{formatted_text}</p>')
    
    for table in doc.tables:
        html_parts.append('<table border="1" style="width: 100%; border-collapse: collapse; margin: 1em 0;">')
        for row in table.rows:
            html_parts.append('<tr>')
            for cell in row.cells:
                # Format cell content preserving bold
                cell_html_parts = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        formatted_text = _format_paragraph_with_bold(paragraph)
                        cell_html_parts.append(formatted_text)
                
                cell_content = '<br>'.join(cell_html_parts) if cell_html_parts else ''
                html_parts.append(f'<td style="padding: 8px; border: 1px solid #ddd;">{cell_content}</td>')
            html_parts.append('</tr>')
        html_parts.append('</table>')
    
    return '\n'.join(html_parts)

def _docx_to_text(docx_path):
    """Convert DOCX to plain text for fallback PDF generation."""
    doc = Document(docx_path)
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            text_parts.append(' | '.join(row_text))
    
    return '\n'.join(text_parts)

def _generate_contract_pdf_fallback(content):
    """
    Fallback PDF generation using reportlab for contracts.
    """
    if not reportlab_available:
        raise Exception("No PDF generation available")
    
    buffer = io.BytesIO()
    doc = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    
    # Simple text rendering
    y = height - 50
    doc.setFont("Helvetica", 12)
    
    lines = content.split('\n')
    for line in lines:
        if y < 50:  # New page
            doc.showPage()
            y = height - 50
            doc.setFont("Helvetica", 12)
        
        if line.strip().startswith('#'):
            # Header
            level = len(line) - len(line.lstrip('#'))
            font_size = 16 - level * 2
            doc.setFont("Helvetica-Bold", font_size)
            doc.drawString(50, y, line.lstrip('#').strip())
            y -= font_size + 10
        else:
            # Regular text
            doc.setFont("Helvetica", 12)
            doc.drawString(50, y, line)
            y -= 15
    
    doc.save()
    buffer.seek(0)
    return buffer.getvalue()




if __name__ == '__main__':
    debug_env = os.getenv('FLASK_DEBUG', 'true').lower() in ('1', 'true', 'yes')
    # Escuchar en 0.0.0.0 para permitir acceso externo cuando se necesite.
    # El puerto puede configurarse con la variable de entorno PORT (por defecto 5000).
    # Prefer explicit PORT; default to 5001 in development to avoid macOS AirDrop on 5000, else 5000
    default_port = '5001' if os.getenv('FLASK_ENV') == 'development' else '5000'
    app.run(debug=debug_env, host='0.0.0.0', port=int(os.getenv('PORT', default_port)))
